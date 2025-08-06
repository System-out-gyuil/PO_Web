import json
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
import requests
import logging
import os
import tempfile
from datetime import datetime
from config import OPEN_AI_API_KEY, AWS_S3_ACCESS_KEY, AWS_S3_SECRET_KEY, AWS_S3_BUCKET_NAME, AWS_S3_REGION
from .models import Row, AttributeValue, Attribute, DropdownAttribute
import pdfplumber
import uuid
import time
import subprocess
from PIL import Image
import warnings
import boto3
from botocore.exceptions import ClientError
import hashlib
import asyncio
import concurrent.futures
import threading
from functools import lru_cache
import mmap
import io
import gzip
import pickle
from pathlib import Path
import zipfile
warnings.filterwarnings("ignore", category=UserWarning)

logger = logging.getLogger(__name__)

# 파일 처리 성능 최적화를 위한 캐시
file_text_cache = {}
file_hash_cache = {}
dropdown_cache = {}  # 드롭다운 캐시 추가
MAX_CACHE_SIZE = 200  # 최대 캐시 크기

# 스레드 풀 생성 (파일 처리용) - 워커 수 증가
file_processing_pool = concurrent.futures.ThreadPoolExecutor(max_workers=8)

# OpenAI API 호출을 위한 별도 스레드 풀
openai_pool = concurrent.futures.ThreadPoolExecutor(max_workers=2)

# 캐시 파일 경로
CACHE_DIR = Path(tempfile.gettempdir()) / "file_cache"
CACHE_DIR.mkdir(exist_ok=True)

# 성능 모니터링을 위한 메트릭
performance_metrics = {
    'file_processing_times': [],
    'openai_api_times': [],
    'cache_hit_rates': []
}

# 파일 크기별 처리 전략
FILE_SIZE_THRESHOLDS = {
    'small': 1024 * 1024,      # 1MB
    'medium': 10 * 1024 * 1024, # 10MB
    'large': 50 * 1024 * 1024   # 50MB
}

# OpenAI API 최적화 설정
OPENAI_API_CONFIG = {
    'timeout': 60,
    'max_retries': 3,
    'retry_delay': 1,
    'batch_size': 5  # 동시 처리할 파일 수
}

def get_cached_file_text(file_hash, file_path):
    """파일 해시 기반 텍스트 캐시 조회 (디스크 + 메모리)"""
    # 메모리 캐시 확인
    if file_hash in file_text_cache:
        logger.info(f"메모리 캐시된 파일 텍스트 사용: {file_hash}")
        return file_text_cache[file_hash]
    
    # 디스크 캐시 확인
    cache_file = CACHE_DIR / f"{file_hash}.pkl"
    if cache_file.exists():
        try:
            with open(cache_file, 'rb') as f:
                cached_data = pickle.load(f)
                # 캐시 유효성 검사 (파일 수정 시간 비교)
                if os.path.exists(file_path):
                    file_mtime = os.path.getmtime(file_path)
                    if cached_data.get('mtime') == file_mtime:
                        text = cached_data.get('text', '')
                        # 메모리 캐시에도 저장
                        set_cached_file_text(file_hash, text)
                        logger.info(f"디스크 캐시된 파일 텍스트 사용: {file_hash}")
                        return text
        except Exception as e:
            logger.error(f"디스크 캐시 읽기 실패: {e}")
    
    return None

def set_cached_file_text(file_hash, text):
    """파일 텍스트 캐시 저장 (메모리 + 디스크)"""
    # 메모리 캐시 관리
    if len(file_text_cache) >= MAX_CACHE_SIZE:
        # LRU 방식으로 가장 오래된 항목 제거
        oldest_key = next(iter(file_text_cache))
        del file_text_cache[oldest_key]
    
    file_text_cache[file_hash] = text
    
    # 디스크 캐시 저장
    try:
        cache_file = CACHE_DIR / f"{file_hash}.pkl"
        cache_data = {
            'text': text,
            'mtime': time.time(),
            'created': time.time()
        }
        with open(cache_file, 'wb') as f:
            pickle.dump(cache_data, f)
    except Exception as e:
        logger.error(f"디스크 캐시 저장 실패: {e}")
    
    logger.info(f"파일 텍스트 캐시 저장: {file_hash}")

def calculate_file_hash_fast(file_path):
    """빠른 파일 해시 계산 (메모리 매핑 사용)"""
    try:
        hash_md5 = hashlib.md5()
        file_size = os.path.getsize(file_path)
        
        if file_size == 0:
            return hashlib.md5(b"").hexdigest()
        
        with open(file_path, 'rb') as f:
            if file_size < 1024 * 1024:  # 1MB 미만
                # 작은 파일은 전체 해시
                with mmap.mmap(f.fileno(), 0, access=mmap.ACCESS_READ) as mm:
                    hash_md5.update(mm)
            else:
                # 큰 파일은 샘플링 방식
                with mmap.mmap(f.fileno(), 0, access=mmap.ACCESS_READ) as mm:
                    # 첫 1KB
                    hash_md5.update(mm[:1024])
                    # 중간 부분 (파일 크기의 50% 지점)
                    mid_pos = file_size // 2
                    hash_md5.update(mm[mid_pos:mid_pos + 1024])
                    # 마지막 1KB
                    hash_md5.update(mm[-1024:])
        
        return hash_md5.hexdigest()
    except Exception as e:
        logger.error(f"빠른 파일 해시 계산 실패: {e}")
        return None

def read_file_chunked(file_path, chunk_size=8192):
    """청크 단위로 파일 읽기 (메모리 효율적)"""
    try:
        with open(file_path, 'rb') as f:
            while True:
                chunk = f.read(chunk_size)
                if not chunk:
                    break
                yield chunk
    except Exception as e:
        logger.error(f"청크 파일 읽기 실패: {e}")

def process_file_async(file_info, field_name):
    """비동기 파일 처리"""
    try:
        file_text, file_hash = extract_file_text(field_name, file_info)
        return file_text, file_hash
    except Exception as e:
        logger.error(f"비동기 파일 처리 실패: {e}")
        return None, None

def extract_text_from_file_optimized(file_path):
    """최적화된 파일 텍스트 추출 (메모리 매핑 사용)"""
    print(f"extract_text_from_file_optimized 시작: {file_path}")
    
    if not file_path or not os.path.exists(file_path):
        print(f"파일이 존재하지 않음: {file_path}")
        return ""
    
    # 파일 크기 체크
    file_size = os.path.getsize(file_path)
    print(f"파일 크기: {file_size} bytes")
    
    # 매우 큰 파일은 처리 제한
    if file_size > 50 * 1024 * 1024:  # 50MB
        print(f"파일이 너무 큼: {file_size / (1024*1024):.1f}MB")
        return f"파일이 너무 커서 텍스트 추출을 건너뜁니다. (크기: {file_size / (1024*1024):.1f}MB)"
    
    # 파일 해시 계산
    file_hash = calculate_file_hash_fast(file_path)
    print(f"파일 해시: {file_hash}")
    
    # 캐시 확인
    cached_text = get_cached_file_text(file_hash, file_path)
    if cached_text:
        print(f"캐시된 텍스트 사용: {len(cached_text)} characters")
        return cached_text
    
    print(f"캐시 없음, 새로 처리 시작")
    
    try:
        file_extension = os.path.splitext(file_path)[1].lower()
        print(f"파일 확장자: {file_extension}")
        
        if file_path.endswith(".pdf"):
            print(f"PDF 파일 처리 시작")
            return extract_pdf_text_optimized(file_path, file_hash)
        elif file_path.endswith((".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff", ".tif", ".webp")):
            print(f"이미지 파일 처리 시작")
            return extract_image_text_optimized(file_path, file_hash)
        elif file_path.endswith(".hwp"):
            print(f"HWP 파일 처리 시작")
            return extract_hwp_text_optimized(file_path, file_hash)
        elif file_path.endswith((".docx", ".doc")):
            print(f"DOCX 파일 처리 시작")
            result = extract_docx_text_optimized(file_path, file_hash)
            print(f"DOCX 처리 결과: {len(result) if result else 0} characters")
            return result
        elif file_path.endswith((".txt", ".md", ".markdown", ".rst", ".adoc")):
            print(f"텍스트 파일 처리 시작")
            return extract_text_file_optimized(file_path, file_hash)
        elif file_path.endswith((".csv", ".tsv")):
            print(f"CSV/TSV 파일 처리 시작")
            return extract_csv_text_optimized(file_path, file_hash)
        elif file_path.endswith((".json", ".xml", ".yaml", ".yml")):
            print(f"구조화된 데이터 파일 처리 시작")
            return extract_structured_text_optimized(file_path, file_hash)
        else:
            print(f"지원하지 않는 파일 형식: {file_extension}")
            return ""
    except Exception as e:
        logger.error(f"최적화된 텍스트 추출 실패: {e}")
        print(f"텍스트 추출 중 오류 발생: {e}")
        return ""

def extract_pdf_text_optimized(file_path, file_hash):
    """최적화된 PDF 텍스트 추출 (메모리 효율적)"""
    try:
        with pdfplumber.open(file_path) as pdf:
            # 페이지 수에 따른 처리 방식 변경
            num_pages = len(pdf.pages)
            
            if num_pages > 50:  # 50페이지 이상은 첫 10페이지만 처리
                pages_to_process = pdf.pages[:10]
                result = "대용량 PDF - 첫 10페이지만 처리됨\n\n"
            else:
                pages_to_process = pdf.pages
                result = ""
            
            all_tables = []
            text_content = []
            
            for i, page in enumerate(pages_to_process):
                # 표 우선 추출
                tables = page.extract_tables()
                if tables:
                    for table in tables:
                        all_tables.append(table)
                
                # 텍스트 추출
                page_text = page.extract_text()
                if page_text:
                    text_content.append(page_text)
            
            # 표가 있으면 표 우선, 없으면 텍스트
            if all_tables:
                table_texts = []
                for table in all_tables:
                    table_texts.append('\n'.join(['\t'.join([cell if cell is not None else '' for cell in row]) for row in table if row]))
                result += '\n\n'.join(table_texts)
            else:
                result += '\n'.join(text_content)
            
            # 캐시 저장
            set_cached_file_text(file_hash, result)
            return result
            
    except Exception as e:
        logger.error(f"PDF 텍스트 추출 실패: {e}")
        return ""

def extract_image_text_optimized(file_path, file_hash):
    """최적화된 이미지 텍스트 추출 (메모리 효율적)"""
    print(f"이미지 텍스트 추출 시작: {file_path}")
    
    try:
        # 이미지 크기 체크 및 리사이즈
        with Image.open(file_path) as img:
            width, height = img.size
            print(f"원본 이미지 크기: {width} x {height}")
            
            if width * height > 4000 * 3000:  # 너무 큰 이미지는 리사이즈
                print(f"이미지가 너무 큼 ({width}x{height}), 리사이즈 진행")
                img.thumbnail((4000, 3000), Image.Resampling.LANCZOS)
                temp_path = tempfile.mktemp(suffix='.jpg')
                img.save(temp_path, 'JPEG', optimize=True, quality=85)
                print(f"리사이즈된 임시 파일 생성: {temp_path}")
                result = clova_ocr(temp_path, 'jpg')
                os.remove(temp_path)
                print(f"임시 파일 삭제 완료")
            else:
                print(f"적정 크기 이미지, 직접 처리")
                # 파일 확장자에 따라 형식 결정
                file_ext = os.path.splitext(file_path)[1].lower()
                if file_ext in ['.jpg', '.jpeg']:
                    fmt = 'jpg'
                elif file_ext == '.png':
                    fmt = 'png'
                elif file_ext == '.gif':
                    fmt = 'gif'
                elif file_ext in ['.bmp']:
                    fmt = 'bmp'
                elif file_ext in ['.tiff', '.tif']:
                    fmt = 'tiff'
                elif file_ext == '.webp':
                    fmt = 'webp'
                else:
                    fmt = 'jpg'  # 기본값
                
                print(f"이미지 형식: {fmt}")
                result = clova_ocr(file_path, fmt)
        
        # 캐시 저장
        if result:
            set_cached_file_text(file_hash, result)
            print(f"이미지 텍스트 추출 완료: {len(result)} characters")
            print(f"추출된 텍스트 내용 (처음 500자): {result[:500]}")
        else:
            print(f"이미지에서 텍스트를 추출하지 못했습니다")
        
        return result
        
    except Exception as e:
        print(f"이미지 텍스트 추출 실패: {e}")
        logger.error(f"이미지 텍스트 추출 실패: {e}")
        return ""

def extract_hwp_text_optimized(file_path, file_hash):
    """최적화된 HWP 텍스트 추출"""
    try:
        # LibreOffice 상태 확인
        if not check_libreoffice_status():
            logger.error("LibreOffice가 설치되지 않았거나 실행할 수 없습니다.")
            return "파일 변환에 실패했습니다."
        
        pdf_path = convert_hwp_to_pdf(file_path)
        if os.path.exists(pdf_path):
            try:
                result = extract_text_from_file_optimized(pdf_path)
                # 변환된 PDF 파일 정리
                if os.path.exists(pdf_path):
                    os.remove(pdf_path)
                    logger.info(f"✅ 임시 PDF 파일 정리 완료: {pdf_path}")
                return result
            except Exception as e:
                logger.error(f"PDF 텍스트 추출 실패: {e}")
                # 변환된 PDF 파일 정리
                if os.path.exists(pdf_path):
                    os.remove(pdf_path)
                return "PDF 텍스트 추출 실패"
        else:
            logger.error("HWP 파일 변환 실패")
            return "HWP 파일 변환 실패"
    except Exception as e:
        logger.error(f"HWP 텍스트 추출 실패: {e}")
        return f"HWP 텍스트 추출 실패: {str(e)}"

def extract_docx_text_optimized(file_path, file_hash):
    """최적화된 DOCX 텍스트 추출"""
    try:
        from docx import Document
        
        # DOCX 파일 열기
        doc = Document(file_path)
        
        # 모든 단락의 텍스트 추출
        paragraphs = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                paragraphs.append(paragraph.text.strip())
        
        # 모든 테이블의 텍스트 추출
        tables = []
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    table_text.append(' | '.join(row_text))
            if table_text:
                tables.append('\n'.join(table_text))
        
        # 결과 조합
        result = []
        if paragraphs:
            result.append('\n'.join(paragraphs))
        if tables:
            result.append('\n\n'.join(tables))
        
        final_result = '\n\n'.join(result) if result else "DOCX 파일에서 텍스트를 추출할 수 없습니다."
        
        # 캐시 저장
        set_cached_file_text(file_hash, final_result)
        return final_result
        
    except ImportError:
        logger.error("python-docx 라이브러리가 설치되지 않았습니다.")
        return "DOCX 파일 처리를 위해 python-docx 라이브러리가 필요합니다."
    except Exception as e:
        logger.error(f"DOCX 텍스트 추출 실패: {e}")
        return f"DOCX 텍스트 추출 실패: {str(e)}"

def extract_text_file_optimized(file_path, file_hash):
    """최적화된 텍스트 파일 추출 (메모리 매핑 사용)"""
    try:
        # 파일 크기 체크
        file_size = os.path.getsize(file_path)
        if file_size > 10 * 1024 * 1024:  # 10MB 이상은 첫 1MB만 읽기
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read(1024 * 1024) + "\n\n[파일이 너무 커서 일부만 표시됩니다.]"
        else:
            # 메모리 매핑을 사용한 효율적인 읽기
            content = ""
            encodings = ['utf-8', 'cp949', 'euc-kr', 'latin-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'rb') as f:
                        with mmap.mmap(f.fileno(), 0, access=mmap.ACCESS_READ) as mm:
                            content = mm.read().decode(encoding)
                    break
                except (UnicodeDecodeError, Exception):
                    continue
            
            if not content:
                return "텍스트 파일 인코딩을 확인할 수 없습니다."
        
        # 캐시 저장
        set_cached_file_text(file_hash, content.strip())
        return content.strip()
        
    except Exception as e:
        logger.error(f"텍스트 파일 추출 실패: {e}")
        return f"텍스트 파일 읽기 실패: {str(e)}"

def extract_csv_text_optimized(file_path, file_hash):
    """최적화된 CSV/TSV 텍스트 추출 (스트리밍 처리)"""
    try:
        import csv
        
        delimiter = ',' if file_path.endswith('.csv') else '\t'
        result = ""
        row_count = 0
        max_rows = 1000
        
        with open(file_path, 'r', encoding='utf-8') as f:
            reader = csv.reader(f, delimiter=delimiter)
            
            for i, row in enumerate(reader):
                if i == 0:  # 헤더
                    result += f"헤더: {' | '.join(row)}\n"
                else:  # 데이터
                    result += f"행 {i}: {' | '.join(row)}\n"
                
                row_count += 1
                if row_count >= max_rows:
                    result += f"\n[대용량 파일 - 첫 {max_rows}행만 처리됨]\n"
                    break
        
        if not result.strip():
            result = "빈 CSV/TSV 파일"
        
        # 캐시 저장
        set_cached_file_text(file_hash, result)
        return result
        
    except Exception as e:
        logger.error(f"CSV/TSV 텍스트 추출 실패: {e}")
        return f"CSV/TSV 파일 읽기 실패: {str(e)}"

def extract_structured_text_optimized(file_path, file_hash):
    """최적화된 구조화된 데이터 파일 추출"""
    try:
        file_size = os.path.getsize(file_path)
        if file_size > 5 * 1024 * 1024:  # 5MB 이상은 일부만 읽기
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read(1024 * 1024) + "\n\n[파일이 너무 커서 일부만 표시됩니다.]"
        else:
            # 메모리 매핑을 사용한 효율적인 읽기
            with open(file_path, 'rb') as f:
                with mmap.mmap(f.fileno(), 0, access=mmap.ACCESS_READ) as mm:
                    content = mm.read().decode('utf-8')
        
        result = f"구조화된 데이터 파일 내용:\n{content.strip()}"
        
        # 캐시 저장
        set_cached_file_text(file_hash, result)
        return result
        
    except Exception as e:
        logger.error(f"구조화된 데이터 파일 추출 실패: {e}")
        return f"구조화된 데이터 파일 읽기 실패: {str(e)}"

def download_file_from_s3_optimized(s3_key):
    """최적화된 S3 파일 다운로드 (스트리밍 + 캐싱)"""
    try:
        # 캐시된 파일 경로 확인
        cache_key = f"s3_download_{hashlib.md5(s3_key.encode()).hexdigest()}"
        cached_path = file_hash_cache.get(cache_key)
        
        if cached_path and os.path.exists(cached_path):
            file_age = time.time() - os.path.getmtime(cached_path)
            if file_age < 18000:  # 캐시 유지 5시간
                logger.info(f"S3 파일 캐시 사용: {s3_key}")
                return cached_path
        
        # S3 클라이언트 생성 (연결 풀링 + 최적화)
        s3_client = boto3.client(
            's3',
            aws_access_key_id=AWS_S3_ACCESS_KEY,
            aws_secret_access_key=AWS_S3_SECRET_KEY,
            region_name=AWS_S3_REGION,
            config=boto3.session.Config(
                max_pool_connections=50,
                retries={'max_attempts': 3},
                read_timeout=30,
                connect_timeout=10
            )
        )
        
        # 임시 파일 생성
        temp_dir = tempfile.gettempdir()
        file_name = s3_key.split('/')[-1]
        temp_path = os.path.join(temp_dir, f"s3_{file_name}")
        
        # S3에서 파일 다운로드 (스트리밍)
        s3_client.download_file(AWS_S3_BUCKET_NAME, s3_key, temp_path)
        
        # 캐시 저장
        file_hash_cache[cache_key] = temp_path
        
        return temp_path
        
    except Exception as e:
        logger.error(f"최적화된 S3 파일 다운로드 실패: {e}")
        return None

def download_file_from_url_optimized(url):
    """최적화된 URL 파일 다운로드 (스트리밍 + 캐싱)"""
    try:
        # S3 URL인 경우 최적화된 S3 다운로드 사용
        if 's3.ap-northeast-2.amazonaws.com' in url:
            return download_file_from_s3_optimized(url)
        
        # 캐시된 파일 경로 확인
        cache_key = f"url_download_{hashlib.md5(url.encode()).hexdigest()}"
        cached_path = file_hash_cache.get(cache_key)
        
        if cached_path and os.path.exists(cached_path):
            # 캐시된 파일이 유효한지 확인 (5분 이내)
            file_age = time.time() - os.path.getmtime(cached_path)
            if file_age < 300:  # 5분
                logger.info(f"URL 파일 캐시 사용: {url}")
                return cached_path
        
        # 일반 URL인 경우 requests 사용 (스트리밍 + 최적화)
        temp_dir = tempfile.gettempdir()
        file_name = url.split('/')[-1].split('?')[0]
        temp_path = os.path.join(temp_dir, f"url_{file_name}")
        
        # 파일 다운로드 (스트리밍, 타임아웃 단축, 청크 크기 최적화)
        response = requests.get(url, stream=True, timeout=15)
        response.raise_for_status()
        
        with open(temp_path, 'wb') as f:
            for chunk in response.iter_content(chunk_size=8192):
                f.write(chunk)
        
        # 캐시 저장
        file_hash_cache[cache_key] = temp_path
        
        return temp_path
        
    except Exception as e:
        logger.error(f"최적화된 URL 파일 다운로드 실패: {e}")
        return None

@csrf_exempt
def ai_chat(request):
    """AI 채팅 API 엔드포인트"""
    if request.method != 'POST':
        return JsonResponse({'success': False, 'error': 'POST 요청만 지원합니다'})
    
    try:
        # JSON 데이터 파싱
        data = json.loads(request.body)
        message = data.get('message', '').strip()
        row_id = data.get('row_id')  # 행 ID 추가
        force_refresh = data.get('force_refresh', False)  # 강제 새로고침 플래그
        changes = data.get('changes', {})  # 변경사항 정보
        
        if not message:
            return JsonResponse({'success': False, 'error': '메시지가 비어있습니다'})
        
        # OpenAI API 키 확인
        if not OPEN_AI_API_KEY:
            return JsonResponse({'success': False, 'error': 'OpenAI API 키가 설정되지 않았습니다'})
        
        # 행 데이터 가져오기 (캐싱 활용)
        row_data = {}
        file_texts = []
        cache_updated = False
        
        if row_id:
            # 세션에서 캐시된 데이터 확인
            cache_key = f'ai_chat_row_{row_id}'
            cached_data = request.session.get(cache_key)
            
            # 현재 시간 정의 (캐시 만료 체크용)
            current_time = time.time()
            
            # 변경사항이 있는지 확인
            has_changes = changes.get('cacheInvalidated', False) or \
                         changes.get('fileChanges', {}).get('added', []) or \
                         changes.get('fileChanges', {}).get('modified', []) or \
                         changes.get('fileChanges', {}).get('deleted', [])
            
            print(f"변경사항 감지 결과:")
            print(f"  cacheInvalidated: {changes.get('cacheInvalidated', False)}")
            print(f"  added files: {len(changes.get('fileChanges', {}).get('added', []))}")
            print(f"  modified files: {len(changes.get('fileChanges', {}).get('modified', []))}")
            print(f"  deleted files: {len(changes.get('fileChanges', {}).get('deleted', []))}")
            print(f"  has_changes: {has_changes}")
            
            # 삭제된 파일 상세 정보 로깅
            deleted_files = changes.get('fileChanges', {}).get('deleted', [])
            if deleted_files:
                print(f"삭제된 파일 상세 정보:")
                for i, deleted_file in enumerate(deleted_files):
                    print(f"  삭제된 파일 {i + 1}: {deleted_file}")
            
            # 변경사항이 있거나 강제 새로고침이면 캐시 무시
            if force_refresh:
                print(f"강제 새로고침으로 인한 새로운 데이터 처리: 행 {row_id}")
                cache_updated = True
                cached_data = None  # 캐시 무시
            elif has_changes and cached_data:
                # 캐시가 있고 변경사항이 있는 경우 - 부분 업데이트만 수행
                print(f"캐시 기반 부분 업데이트: 행 {row_id}")
                cache_updated = True
                # cached_data는 유지 (None으로 설정하지 않음)
            elif has_changes and not cached_data:
                # 캐시가 없고 변경사항이 있는 경우 - 새로운 데이터 처리
                print(f"캐시 없음 + 변경사항으로 인한 새로운 데이터 처리: 행 {row_id}")
                cache_updated = True
                cached_data = None
            elif cached_data:
                # 캐시 만료 시간 체크 (30분)
                cache_timestamp = cached_data.get('timestamp', 0)
                cache_age = current_time - cache_timestamp
                
                if cache_age < 1800:  # 30분 (1800초)
                    # 캐시된 데이터 사용
                    row_data = cached_data.get('row_data', {})
                    file_texts = cached_data.get('file_texts', [])
                    print(f"캐시된 데이터 사용: 행 {row_id}")
                    
                    # text 타입 파일들은 항상 최신 데이터로 업데이트
                    try:
                        row = Row.objects.get(id=row_id)
                        text_file_attribute_values = AttributeValue.objects.filter(row=row, attribute__attributeType__name='file')
                        
                        for attr_value in text_file_attribute_values:
                            attr_name = attr_value.attribute.name
                            
                            if attr_value.value:
                                try:
                                    file_data = json.loads(attr_value.value) if isinstance(attr_value.value, str) else attr_value.value
                                    print(f'캐시 사용 중 text 타입 확인 - file_data: {file_data}')
                                    
                                    # 음성파일 속성인 경우 (data 구조)
                                    if isinstance(file_data, dict) and 'data' in file_data:
                                        for file_id, file_info in file_data['data'].items():
                                            if file_info.get('type') == 'text':
                                                # text 타입은 항상 새로 처리
                                                print(f'캐시 사용 중 text 타입 새로 처리: {attr_name}')
                                                text_content = file_info.get('text', '')
                                                if text_content:
                                                    # 기존 text 타입 파일 제거
                                                    file_texts = [text for text in file_texts if not text.startswith(f"[{attr_name} - 텍스트]:")]
                                                    # 새 text 내용 추가
                                                    file_texts.append(f"[{attr_name} - 텍스트]:\n{text_content}")
                                                    print(f'text 타입 파일 업데이트 완료: {attr_name}')
                                    
                                    # 일반 파일 속성인 경우 (배열 구조) - text 타입 처리
                                    elif isinstance(file_data, list):
                                        for file_info in file_data:
                                            if file_info.get('type') == 'text':
                                                print(f'캐시 사용 중 text 타입 새로 처리: {attr_name}')
                                                text_content = file_info.get('text', '')
                                                if text_content:
                                                    # 기존 text 타입 파일 제거
                                                    file_texts = [text for text in file_texts if not text.startswith(f"[{attr_name} - 텍스트]:")]
                                                    # 새 text 내용 추가
                                                    file_texts.append(f"[{attr_name} - 텍스트]:\n{text_content}")
                                                    print(f'text 타입 파일 업데이트 완료: {attr_name}')
                                except Exception as e:
                                    logger.error(f"캐시 사용 중 text 타입 파일 처리 실패: {e}")
                    except Exception as e:
                        logger.error(f"캐시 사용 중 text 타입 파일 새로 처리 중 오류: {e}")
                    
                    # 업데이트된 file_texts로 캐시 갱신
                    cache_data = {
                        'row_data': row_data,
                        'file_texts': file_texts,
                        'timestamp': current_time
                    }
                    request.session[cache_key] = cache_data
                    request.session.modified = True
                    cache_updated = True
                    print(f"text 타입 파일 업데이트로 인한 캐시 갱신 완료")
                else:
                    # 캐시 만료 - 새로운 데이터 처리
                    print(f"캐시 만료 - 새로운 데이터 처리: 행 {row_id}")
                    cache_updated = True
                    cached_data = None  # 캐시 무시
            else:
                # 캐시가 없음 - 새로운 데이터 처리
                print(f"캐시 없음 - 새로운 데이터 처리: 행 {row_id}")
                cache_updated = True
            
            # 새로운 데이터 처리 (캐시가 없거나 무효화된 경우)
            if not cached_data:
                try:
                    row = Row.objects.get(id=row_id)
                    user = row.user
                    
                    # 해당 행의 모든 속성값 가져오기 (파일 제외)
                    attribute_values = AttributeValue.objects.filter(row=row)
                    
                    for attr_value in attribute_values:
                        attr_name = attr_value.attribute.name
                        attr_type = attr_value.attribute.attributeType.name
                        
                        # 파일 타입은 건너뛰고 나머지 데이터만 처리
                        if attr_type != 'file':
                            if attr_type == 'outstanding_debts':
                                # 기대출 데이터 처리
                                if attr_value.value:
                                    try:
                                        debt_data = json.loads(attr_value.value) if isinstance(attr_value.value, str) else attr_value.value
                                        if isinstance(debt_data, dict):
                                            # 영어 키를 한글로 변환하는 매핑
                                            key_mapping = {
                                                'credit_foundation': '신용보증재단',
                                                'credit_guarantee': '신용보증',
                                                'sbc': 'SBC',
                                                'kibo': '기보',
                                                'kibo_foundation': '기보재단',
                                                'kibo_guarantee': '기보보증',
                                                'sbc_foundation': 'SBC재단',
                                                'sbc_guarantee': 'SBC보증',
                                                'other': '기타',
                                                'total': '총합',
                                                'total_amount': '총금액',
                                                'amount': '금액',
                                                'limit': '한도',
                                                'used': '사용액',
                                                'available': '가용액',
                                                'remaining': '잔여액'
                                            }
                                            
                                            debt_summary = []
                                            for key, value in debt_data.items():
                                                if value and value != 0:
                                                    # 영어 키를 한글로 변환
                                                    korean_key = key_mapping.get(key, key)
                                                    debt_summary.append(f"{korean_key}: {value:,}만원")
                                            if debt_summary:
                                                row_data[attr_name] = " | ".join(debt_summary)
                                    except Exception as e:
                                        logger.error(f"기대출 데이터 처리 실패: {e}")
                            
                            elif attr_type == 'recommend':
                                # 추천 데이터 처리
                                if attr_value.value:
                                    try:
                                        recommend_data = json.loads(attr_value.value) if isinstance(attr_value.value, str) else attr_value.value
                                        if isinstance(recommend_data, dict):
                                            # 영어 키를 한글로 변환하는 매핑
                                            recommend_key_mapping = {
                                                'fund_name': '자금명',
                                                'limit': '한도',
                                                'institution': '기관',
                                                'amount': '금액',
                                                'total': '총합',
                                                'total_amount': '총금액',
                                                'available': '가용액',
                                                'used': '사용액',
                                                'remaining': '잔여액',
                                                'interest_rate': '이자율',
                                                'period': '기간',
                                                'type': '유형',
                                                'category': '카테고리'
                                            }
                                            
                                            # 총 자금 정보
                                            total_funds = recommend_data.get('총자금', 0)
                                            if total_funds:
                                                row_data[f"{attr_name}_총자금"] = f"{total_funds:,}원"
                                            
                                            # 자금들 정보
                                            funds = recommend_data.get('자금들', {})
                                            if funds:
                                                fund_summary = []
                                                for fund_name, amount in funds.items():
                                                    if amount and amount > 0:
                                                        fund_summary.append(f"{fund_name}: {amount:,}원")
                                                if fund_summary:
                                                    row_data[f"{attr_name}_자금들"] = " | ".join(fund_summary)
                                            
                                            # 상세정보
                                            detail_info = recommend_data.get('상세정보', [])
                                            if detail_info:
                                                detail_summary = []
                                                for detail in detail_info:
                                                    fund_name = detail.get('fund_name', '')
                                                    limit = detail.get('limit', 0)
                                                    institution = detail.get('institution', '')
                                                    
                                                    # 영어 키를 한글로 변환
                                                    korean_fund_name = recommend_key_mapping.get('fund_name', fund_name)
                                                    korean_limit = recommend_key_mapping.get('limit', '한도')
                                                    korean_institution = recommend_key_mapping.get('institution', '기관')
                                                    
                                                    if fund_name and limit:
                                                        detail_summary.append(f"{fund_name}({institution}): {limit:,}원")
                                                if detail_summary:
                                                    row_data[f"{attr_name}_상세정보"] = " | ".join(detail_summary)
                                    except Exception as e:
                                        logger.error(f"추천 데이터 처리 실패: {e}")
                            
                            elif attr_type == 'dropdown':
                                # 드롭다운 데이터 처리
                                if attr_value.value:
                                    try:
                                        dropdown_value = str(attr_value.value).strip()
                                        if dropdown_value:
                                            dropdown_name = get_dropdown_name(attr_value.attribute, dropdown_value)
                                            row_data[attr_name] = dropdown_name
                                            print(f"캐시 기반 드롭다운 업데이트: {attr_name} = {dropdown_name}")
                                    except Exception as e:
                                        logger.error(f"드롭다운 데이터 처리 실패: {e}")
                                        row_data[attr_name] = f"{attr_value.value} (코드의 의미는 데이터에 없음)"
                            
                            elif attr_type == 'text':
                                # 텍스트 데이터 처리
                                if attr_value.value:
                                    row_data[attr_name] = attr_value.value
                            
                            elif attr_type == 'number':
                                # 숫자 데이터 처리
                                if attr_value.value:
                                    try:
                                        num_value = float(attr_value.value)
                                        row_data[attr_name] = f"{num_value:,}"
                                    except:
                                        row_data[attr_name] = attr_value.value
                            
                            elif attr_type == 'date':
                                # 날짜 데이터 처리
                                if attr_value.value:
                                    row_data[attr_name] = attr_value.value
                            
                            elif attr_type == 'boolean':
                                # 불린 데이터 처리
                                if attr_value.value:
                                    bool_value = attr_value.value.lower() if isinstance(attr_value.value, str) else str(attr_value.value)
                                    if bool_value in ['true', '1', 'yes', 'on']:
                                        row_data[attr_name] = "예"
                                    elif bool_value in ['false', '0', 'no', 'off']:
                                        row_data[attr_name] = "아니오"
                                    else:
                                        row_data[attr_name] = attr_value.value
                            
                            else:
                                # 기타 타입
                                if attr_value.value:
                                    row_data[attr_name] = attr_value.value
                    
                    # 파일 데이터 처리 - 최적화된 방식으로 교체
                    file_texts = []
                    
                    # 파일 속성값들 처리
                    file_attribute_values = AttributeValue.objects.filter(row=row, attribute__attributeType__name='file')
                    
                    # 모든 파일 정보 수집
                    all_files = []
                    
                    for attr_value in file_attribute_values:
                        attr_name = attr_value.attribute.name
                        print(f"=== 파일 속성 처리: {attr_name} ===")
                        
                        if attr_value.value:
                            try:
                                # 파일 데이터 파싱
                                file_data = json.loads(attr_value.value) if isinstance(attr_value.value, str) else attr_value.value
                                print(f"파일 데이터 타입: {type(file_data)}")
                                print(f"파일 데이터 구조: {file_data}")
                                
                                # 음성파일 속성인 경우 (data 구조)
                                if isinstance(file_data, dict) and 'data' in file_data:
                                    print(f"data 구조 파일 개수: {len(file_data['data'])}")
                                    for file_id, file_info in file_data['data'].items():
                                        print(f"파일 ID: {file_id}, 파일 정보: {file_info}")
                                        if file_info.get('type') in ['file', 'image']:  # 'image' 타입도 추가
                                            filename = file_info.get('original_filename', '파일')
                                            print(f"파일 감지 (data 구조): {filename} (속성: {attr_name}) - 타입: {file_info.get('type')}")
                                            
                                            # 파일 크기 체크 (30MB 제한)
                                            file_size = file_info.get('file_size', 0)
                                            if file_size > 30 * 1024 * 1024:  # 30MB
                                                file_texts.append(f"[{attr_name} - {file_info.get('original_filename', '파일')}]: 파일이 너무 커서 텍스트 추출을 건너뜁니다.")
                                                print(f"파일 크기 초과로 건너뜀: {filename} ({file_size / (1024*1024):.1f}MB)")
                                                continue
                                            
                                            # 파일 정보에 필드명 추가
                                            file_info['field_name'] = attr_name
                                            all_files.append(file_info)
                                            print(f"파일 처리 대상 추가 (data 구조): {filename}")
                                elif file_info.get('type') == 'text':
                                    # text 타입은 항상 새로 처리 (캐시 무시)
                                    print(f'text 타입 파일 새로 처리: {file_info.get("text")}')
                                    text_content = file_info.get('text', '')
                                    if text_content:
                                        file_texts.append(f"[{attr_name} - 텍스트]:\n{text_content}")
                                        print(f'text 타입 파일 캐시에 저장: {attr_name}')
                                
                                # 일반 파일 속성인 경우 (배열 구조)
                                elif isinstance(file_data, list):
                                    print(f"배열 구조 파일 개수: {len(file_data)}")
                                    for i, file_info in enumerate(file_data):
                                        print(f"파일 {i}: {file_info}")
                                        filename = file_info.get('original_filename', '파일')
                                        print(f"파일 감지 (배열 구조): {filename} (속성: {attr_name})")
                                        
                                        # 파일 크기 체크 (10MB 제한)
                                        file_size = file_info.get('file_size', 0)
                                        if file_size > 10 * 1024 * 1024:  # 10MB
                                            file_texts.append(f"[{attr_name} - {file_info.get('original_filename', '파일')}]: 파일이 너무 커서 텍스트 추출을 건너뜁니다.")
                                            print(f"파일 크기 초과로 건너뜀: {filename} ({file_size / (1024*1024):.1f}MB)")
                                            continue
                                        
                                        # 파일 정보에 필드명 추가
                                        file_info['field_name'] = attr_name
                                        all_files.append(file_info)
                                        print(f"파일 처리 대상 추가 (배열 구조): {filename}")
                                
                                # 단일 파일 경로인 경우
                                elif isinstance(file_data, str):
                                    print(f"단일 파일 경로: {file_data}")
                                    file_path = file_data
                                    if file_path.startswith('http'):
                                        file_path = download_file_from_url_optimized(file_path)
                                    
                                    if file_path and os.path.exists(file_path):
                                        file_text = extract_text_from_file_optimized(file_path)
                                        file_hash = calculate_file_hash_fast(file_path)
                                        if file_text:
                                            file_texts.append(f"[{attr_name} 파일 내용]:\n{file_text}")
                                        
                                        # 임시 파일 정리
                                        if file_data.startswith('http'):
                                            try:
                                                os.remove(file_path)
                                            except:
                                                pass
                                else:
                                    print(f"알 수 없는 파일 데이터 구조: {type(file_data)}")
                                    
                            except Exception as e:
                                logger.error(f"파일 텍스트 추출 실패: {e}")
                                print(f"파일 처리 중 오류: {e}")
                        else:
                            print(f"속성 {attr_name}에 값이 없음")
                    
                    # 최적화된 파일 처리 전략 적용
                    if all_files:
                        print(f"최적화된 파일 처리 시작: {len(all_files)}개 파일")
                        
                        # 파일 크기별 분류
                        file_groups = optimize_file_processing_strategy(all_files)
                        
                        # 우선순위 기반 처리
                        processed_file_texts = process_files_with_priority(file_groups, "optimized_files")
                        file_texts.extend(processed_file_texts)
                        
                        print(f"파일 처리 완료: {len(processed_file_texts)}개 파일 텍스트 추출")
                        
                        # 추출된 텍스트 내용 로깅 (이미지 파일 포함)
                        for i, text in enumerate(processed_file_texts):
                            print(f"파일 텍스트 {i+1}: {text[:200]}...")
                    
                    # 전체 파일 텍스트 결과 로깅
                    print(f"총 파일 텍스트 개수: {len(file_texts)}")
                    for i, text in enumerate(file_texts):
                        if "이미지" in text or ".jpg" in text or ".png" in text or ".gif" in text:
                            print(f"이미지 관련 텍스트 {i+1}: {text[:300]}...")
                    
                except Row.DoesNotExist:
                    return JsonResponse({'success': False, 'error': '해당 행을 찾을 수 없습니다'})
                except Exception as e:
                    logger.error(f"행 데이터 조회 실패: {e}")
            else:
                # 캐시가 있고 변경사항이 있는 경우 - 부분 업데이트만 수행
                if has_changes:
                    print(f"캐시 기반 부분 업데이트: 행 {row_id}")
                    print(f"  캐시 데이터 존재: {cached_data is not None}")
                    print(f"  변경사항 존재: {has_changes}")
                    
                    # 기존 캐시 데이터 사용
                    row_data = cached_data.get('row_data', {})
                    file_texts = cached_data.get('file_texts', [])
                    
                    print(f"  기존 file_texts 개수: {len(file_texts)}")
                    for i, text in enumerate(file_texts):
                        print(f"    기존 file_text {i + 1}: {text[:100]}...")
                    
                    # 드롭다운 데이터 최신 상태로 업데이트
                    try:
                        row = Row.objects.get(id=row_id)
                        dropdown_attribute_values = AttributeValue.objects.filter(row=row, attribute__attributeType__name='dropdown')
                        
                        for attr_value in dropdown_attribute_values:
                            attr_name = attr_value.attribute.name
                            
                            if attr_value.value:
                                try:
                                    dropdown_value = str(attr_value.value).strip()
                                    if dropdown_value:
                                        dropdown_name = get_dropdown_name(attr_value.attribute, dropdown_value)
                                        row_data[attr_name] = dropdown_name
                                        print(f"캐시 기반 드롭다운 업데이트: {attr_name} = {dropdown_name}")
                                except Exception as e:
                                    logger.error(f"캐시 기반 드롭다운 데이터 처리 실패: {e}")
                                    row_data[attr_name] = f"{attr_value.value} (코드의 의미는 데이터에 없음)"
                    except Exception as e:
                        logger.error(f"캐시 기반 드롭다운 업데이트 중 오류: {e}")
                    
                    # text 타입 파일들은 항상 새로 처리 (캐시 무시)
                    try:
                        row = Row.objects.get(id=row_id)
                        text_file_attribute_values = AttributeValue.objects.filter(row=row, attribute__attributeType__name='file')
                        
                        for attr_value in text_file_attribute_values:
                            attr_name = attr_value.attribute.name
                            
                            if attr_value.value:
                                try:
                                    file_data = json.loads(attr_value.value) if isinstance(attr_value.value, str) else attr_value.value
                                    
                                    print(f'file_data: {file_data}')
                                    # 음성파일 속성인 경우 (data 구조)
                                    if isinstance(file_data, dict) and 'data' in file_data:
                                        for file_id, file_info in file_data['data'].items():
                                            if file_info.get('type') == 'text':
                                                # text 타입은 항상 새로 처리
                                                print(f'캐시 기반 부분 업데이트에서 text 타입 새로 처리: {attr_name}')
                                                text_content = file_info.get('text', '')
                                                if text_content:
                                                    # 기존 text 타입 파일 제거
                                                    file_texts = [text for text in file_texts if not text.startswith(f"[{attr_name} - 텍스트]:")]
                                                    # 새 text 내용 추가
                                                    file_texts.append(f"[{attr_name} - 텍스트]:\n{text_content}")
                                                    print(f'text 타입 파일 업데이트 완료: {attr_name}')
                                    
                                    # 일반 파일 속성인 경우 (배열 구조) - text 타입 처리
                                    elif isinstance(file_data, list):
                                        for file_info in file_data:
                                            if file_info.get('type') == 'text':
                                                print(f'캐시 기반 부분 업데이트에서 text 타입 새로 처리: {attr_name}')
                                                text_content = file_info.get('text', '')
                                                if text_content:
                                                    # 기존 text 타입 파일 제거
                                                    file_texts = [text for text in file_texts if not text.startswith(f"[{attr_name} - 텍스트]:")]
                                                    # 새 text 내용 추가
                                                    file_texts.append(f"[{attr_name} - 텍스트]:\n{text_content}")
                                                    print(f'text 타입 파일 업데이트 완료: {attr_name}')
                                except Exception as e:
                                    logger.error(f"캐시 기반 text 타입 파일 처리 실패: {e}")
                    except Exception as e:
                        logger.error(f"text 타입 파일 새로 처리 중 오류: {e}")
                    
                    # 파일 변경사항 처리
                    file_changes = changes.get('fileChanges', {})
                    
                    # 삭제된 파일 처리
                    deleted_files = file_changes.get('deleted', [])
                    for deleted_file in deleted_files:
                        field_name = deleted_file.get('fieldName')
                        file_info = deleted_file.get('fileInfo', {})
                        if field_name:
                            print(f"파일 삭제 처리 시작: {field_name}")
                            print(f"  삭제된 파일 정보: {file_info}")
                            
                            original_file_texts_count = len(file_texts)
                            
                            # 삭제할 파일의 식별자들
                            target_filename = file_info.get('original_filename')
                            target_file_hash = file_info.get('file_hash')
                            target_file_id = file_info.get('fileId')
                            target_s3_key = file_info.get('s3_key')
                            
                            print(f"  삭제 대상 파일 식별자:")
                            print(f"    파일명: {target_filename}")
                            print(f"    파일해시: {target_file_hash}")
                            print(f"    파일ID: {target_file_id}")
                            print(f"    S3키: {target_s3_key}")
                            
                            print(f"  기존 캐시 파일들 ({len(file_texts)}개):")
                            for i, text in enumerate(file_texts):
                                print(f"    캐시 파일 {i+1}: {text[:100]}...")
                            
                            # 정확한 파일 매칭을 위한 필터링 함수
                            def should_keep_file_text(text):
                                # text 타입 파일은 항상 유지 (삭제 대상 아님)
                                if text.startswith(f"[{field_name} - 텍스트]:"):
                                    return True
                                
                                # 삭제 대상 파일인지 확인
                                is_target_file = False
                                
                                # 1. 필드명 + 파일명으로 매칭 (정확한 매칭)
                                if target_filename and f"[{field_name} - {target_filename}]:" in text:
                                    is_target_file = True
                                    print(f"    필드명+파일명 매칭으로 삭제 대상 확인: {field_name} - {target_filename}")
                                    print(f"    매칭된 텍스트: {text[:200]}...")
                                
                                # 2. 파일명만으로 매칭 (필드명과 관계없이)
                                elif target_filename and f" - {target_filename}]:" in text:
                                    is_target_file = True
                                    print(f"    파일명만 매칭으로 삭제 대상 확인: {target_filename}")
                                    print(f"    매칭된 텍스트: {text[:200]}...")
                                
                                # 3. fileId로 매칭 (음성파일의 경우)
                                elif target_file_id and f"fileId: {target_file_id}" in text:
                                    is_target_file = True
                                    print(f"    fileId 매칭으로 삭제 대상 확인: {target_file_id}")
                                    print(f"    매칭된 텍스트: {text[:200]}...")
                                
                                # 4. s3_key로 매칭
                                elif target_s3_key and target_s3_key in text:
                                    is_target_file = True
                                    print(f"    s3_key 매칭으로 삭제 대상 확인: {target_s3_key}")
                                    print(f"    매칭된 텍스트: {text[:200]}...")
                                
                                # 5. 해시로 매칭 (정확한 해시 정보가 있는 경우)
                                elif target_file_hash and f"해시: {target_file_hash}" in text:
                                    is_target_file = True
                                    print(f"    해시 매칭으로 삭제 대상 확인: {target_file_hash}")
                                    print(f"    매칭된 텍스트: {text[:200]}...")
                                
                                # 6. 파일 경로에서 파일명 추출하여 매칭
                                elif target_filename:
                                    # 파일 경로에서 파일명만 추출하여 비교
                                    import os
                                    extracted_filename = os.path.basename(target_filename) if '/' in target_filename else target_filename
                                    if f" - {extracted_filename}]:" in text:
                                        is_target_file = True
                                        print(f"    추출된 파일명 매칭으로 삭제 대상 확인: {extracted_filename}")
                                        print(f"    매칭된 텍스트: {text[:200]}...")
                                
                                # 삭제 대상이면 False 반환 (제거), 아니면 True 반환 (유지)
                                if is_target_file:
                                    print(f"    삭제 대상 파일 확인됨 - 제거 예정")
                                return not is_target_file
                            
                            # 파일 텍스트 필터링
                            file_texts = [text for text in file_texts if should_keep_file_text(text)]
                            
                            # 삭제된 파일의 캐시도 정리
                            if target_file_hash:
                                # 메모리 캐시에서 제거
                                if target_file_hash in file_text_cache:
                                    del file_text_cache[target_file_hash]
                                    print(f"    메모리 캐시에서 파일 해시 제거: {target_file_hash}")
                                
                                # 디스크 캐시에서 제거
                                cache_file = CACHE_DIR / f"{target_file_hash}.pkl"
                                if cache_file.exists():
                                    try:
                                        cache_file.unlink()
                                        print(f"    디스크 캐시에서 파일 해시 제거: {target_file_hash}")
                                    except Exception as e:
                                        print(f"    디스크 캐시 제거 실패: {e}")
                            
                            # s3_key 기반 캐시도 정리
                            if target_s3_key:
                                cache_key = f"s3_download_{hashlib.md5(target_s3_key.encode()).hexdigest()}"
                                if cache_key in file_hash_cache:
                                    cached_path = file_hash_cache[cache_key]
                                    if os.path.exists(cached_path):
                                        try:
                                            os.remove(cached_path)
                                            print(f"    S3 캐시 파일 제거: {cached_path}")
                                        except Exception as e:
                                            print(f"    S3 캐시 파일 제거 실패: {e}")
                                    del file_hash_cache[cache_key]
                                    print(f"    S3 캐시 키 제거: {cache_key}")
                            
                            removed_count = original_file_texts_count - len(file_texts)
                            print(f"  제거된 파일 텍스트 수: {removed_count}")
                            print(f"파일 삭제 처리 완료: {field_name}")
                    
                    # 추가된 파일들 처리
                    added_files = file_changes.get('added', [])
                    for added_file in added_files:
                        field_name = added_file.get('fieldName')
                        file_info = added_file.get('fileInfo', {})
                        if field_name and file_info:
                            try:
                                print(f"파일 추가 처리 시작: {field_name}")
                                print(f"  파일명: {file_info.get('original_filename', 'N/A')}")
                                print(f"  해시: {file_info.get('file_hash', 'N/A')}")
                                print(f"  크기: {file_info.get('file_size', 'N/A')}")
                                
                                # 새 파일 텍스트 추출
                                file_text, file_hash = extract_file_text(field_name, file_info)
                                if file_text:
                                    file_texts.append(file_text)
                                    print(f"파일 추가 처리: {field_name} - {file_info.get('original_filename', '파일')} (해시: {file_hash or 'N/A'})")
                                else:
                                    print(f"파일 추가 처리 실패: 텍스트 추출 실패")
                            except Exception as e:
                                logger.error(f"추가된 파일 텍스트 추출 실패: {e}")
                                print(f"파일 추가 처리 중 오류: {e}")
                    
                    # 수정된 파일들 처리
                    modified_files = file_changes.get('modified', [])
                    for modified_file in modified_files:
                        field_name = modified_file.get('fieldName')
                        file_info = modified_file.get('fileInfo', {})
                        if field_name and file_info:
                            try:
                                print(f"파일 수정 처리 시작: {field_name}")
                                print(f"  파일명: {file_info.get('original_filename', 'N/A')}")
                                print(f"  새 해시: {file_info.get('file_hash', 'N/A')}")
                                
                                # 기존 파일 텍스트 제거 (다양한 매칭 방식 사용)
                                original_file_texts_count = len(file_texts)
                                target_filename = file_info.get('original_filename', '')
                                
                                # 정확한 파일 매칭을 위한 제거 함수
                                def should_remove_file_text(text):
                                    # text 타입 파일은 수정 대상 아님
                                    if text.startswith(f"[{field_name} - 텍스트]:"):
                                        return False
                                    
                                    # 1. 필드명 + 파일명으로 매칭
                                    if target_filename and f"[{field_name} - {target_filename}]:" in text:
                                        return True
                                    
                                    # 2. 파일명만으로 매칭
                                    if target_filename and f" - {target_filename}]:" in text:
                                        return True
                                    
                                    # 3. 파일 경로에서 파일명 추출하여 매칭
                                    if target_filename:
                                        import os
                                        extracted_filename = os.path.basename(target_filename) if '/' in target_filename else target_filename
                                        if f" - {extracted_filename}]:" in text:
                                            return True
                                    
                                    return False
                                
                                # 파일 텍스트 필터링
                                file_texts = [text for text in file_texts if not should_remove_file_text(text)]
                                removed_count = original_file_texts_count - len(file_texts)
                                print(f"  제거된 파일 텍스트 수: {removed_count}")
                                
                                # 새 파일 텍스트 추출
                                file_text, file_hash = extract_file_text(field_name, file_info)
                                if file_text:
                                    file_texts.append(file_text)
                                    print(f"파일 수정 처리: {field_name} - {file_info.get('original_filename', '파일')} (해시: {file_hash or 'N/A'})")
                                else:
                                    print(f"파일 수정 처리 실패: 텍스트 추출 실패")
                            except Exception as e:
                                logger.error(f"수정된 파일 텍스트 추출 실패: {e}")
                                print(f"파일 수정 처리 중 오류: {e}")
                    
                    # 파일 변경사항이 있었으면 캐시 무효화 강화
                    if deleted_files or added_files or modified_files:
                        print(f"파일 변경사항으로 인한 캐시 무효화 강화")
                        # 관련된 모든 캐시 정리
                        cleanup_related_caches(deleted_files, added_files, modified_files)
                    
                    # 캐시 업데이트
                    cache_data = {
                        'row_data': row_data,
                        'file_texts': file_texts,
                        'timestamp': current_time
                    }
                    request.session[cache_key] = cache_data
                    request.session.modified = True
                    cache_updated = True
                    
                    print(f"[AI 캐시 부분 업데이트] row_id={row_id}")
                    print(f"  추가된 파일: {len(added_files)}개")
                    print(f"  수정된 파일: {len(modified_files)}개")
                    print(f"  삭제된 파일: {len(deleted_files)}개")
                    for file_text in file_texts:
                        print(f"  최종 file_text: {file_text}")
        
        # OpenAI API 호출 - 최적화된 방식으로 교체
        # 컨텍스트 최적화
        optimized_context = optimize_context_for_openai(row_data, file_texts)
        
        # 오늘 날짜 정보 추가
        from datetime import datetime, timezone
        today_str = datetime.now(timezone.utc).astimezone().strftime('%Y-%m-%d %H:%M:%S %Z')
        
        # 최적화된 컨텍스트 정보 구성
        context_info = f"\n\n[오늘 날짜(현실 기준)]: {today_str}\n\n"
        if optimized_context:
            context_info += f"\n{optimized_context}\n"
        
        # 영업 관련 컨텍스트를 포함한 프롬프트 생성
        system_prompt = """당신은 정확하고 신뢰할 수 있는 AI 어시스턴트입니다. 
        
        **중요한 지침:**
        1. 제공된 데이터와 파일 내용만을 기반으로 답변하세요.
        2. 데이터에 없는 정보는 추측하거나 지어내지 마세요.
        3. 확실하지 않은 정보는 "확인되지 않음" 또는 "데이터에 없음"이라고 명시하세요.
        4. 답변은 친근하고 실용적이며, 한국어로 작성해주세요.
        
        **영업 관련 전문 분야:**
        - 영업 전략 및 전략 수립
        - 고객 관리 및 리드 관리
        - 매출 증대 및 성과 개선
        - 재무 분석 및 자금 조달
        - 시장 분석 및 경쟁 분석
        
        **데이터 활용 방법:**
        - 회사 기본 정보 (회사명, 업종, 지역, 설립일 등)
        - 재무 정보 (매출, 기대출, 추천 자금, 자금 조달 현황)
        - 첨부된 문서 내용 (재무제표, 대출내역, 확인서, 계약서 등)
        - 날짜 정보 (설립일, F/U 일정, 계약일 등)
        - 기타 관련 정보 (담당자, 연락처, 특이사항 등)
        
        **답변 형식:**
        - 구체적이고 실용적인 조언 제공
        - 데이터 기반의 객관적 분석
        - 필요한 경우 추가 정보 요청
        - 확실하지 않은 부분은 명확히 구분하여 표시"""
        
        user_message = f"{message}{context_info}"
        
        # 최적화된 OpenAI API 호출
        api_result = call_openai_api_optimized([
            {'role': 'system', 'content': system_prompt},
            {'role': 'user', 'content': user_message}
        ])
        
        if api_result['success']:
            return JsonResponse({
                'success': True,
                'response': api_result['response'],
                'cache_updated': cache_updated,
                'processing_time': api_result.get('time', 0)
            })
        else:
            return JsonResponse({'success': False, 'error': api_result['error']})
            
    except requests.exceptions.Timeout:
        return JsonResponse({'success': False, 'error': '요청 시간이 초과되었습니다. 다시 시도해주세요.'})
    except requests.exceptions.RequestException as e:
        return JsonResponse({'success': False, 'error': f'네트워크 오류: {str(e)}'})
    except json.JSONDecodeError:
        return JsonResponse({'success': False, 'error': '잘못된 JSON 형식입니다'})
    except Exception as e:
        logger.error(f"AI 채팅 오류: {str(e)}")
        return JsonResponse({'success': False, 'error': f'서버 오류: {str(e)}'})

def calculate_file_hash(file_path):
    """파일의 MD5 해시를 계산"""
    try:
        hash_md5 = hashlib.md5()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_md5.update(chunk)
        return hash_md5.hexdigest()
    except Exception as e:
        logger.error(f"파일 해시 계산 실패: {e}")
        return None

def extract_file_text(field_name, file_info):
    """파일 정보에서 텍스트 추출하는 헬퍼 함수"""
    try:
        # 파일 크기 체크 (10MB 제한)
        file_size = file_info.get('file_size', 0)
        if file_size > 10 * 1024 * 1024:  # 10MB
            return f"[{field_name} - {file_info.get('original_filename', '파일')}]: 파일이 너무 커서 텍스트 추출을 건너뜁니다.", None
        
        file_path = None
        # S3 키가 있으면 직접 사용
        s3_key = file_info.get('s3_key')
        if s3_key:
            file_path = download_file_from_s3_key(s3_key)
        # S3 키가 없고 download_url이 있으면 사용
        elif file_info.get('download_url'):
            file_path = download_file_from_url(file_info['download_url'])
        
        if file_path and os.path.exists(file_path):
            print(f"파일 처리 시작: {file_path}")
            print(f"파일 크기: {os.path.getsize(file_path)} bytes")
            print(f"파일 확장자: {os.path.splitext(file_path)[1].lower()}")
            
            file_text = extract_text_from_file_optimized(file_path)
            file_hash = calculate_file_hash(file_path)
            
            print(f"텍스트 추출 결과: {len(file_text) if file_text else 0} characters")
            if file_text and len(file_text) > 0:
                print(f"추출된 텍스트 샘플: {file_text[:200]}...")
            
            if file_text:
                result = f"[{field_name} - {file_info.get('original_filename', '파일')}]:\n{file_text}"
                # 파일 해시 정보를 로그에 추가
                if file_hash:
                    print(f"파일 해시 계산 완료: {file_hash} - {file_info.get('original_filename', '파일')}")
                return result, file_hash
            else:
                result = f"[{field_name} - {file_info.get('original_filename', '파일')}]: 텍스트 추출 실패"
                if file_hash:
                    print(f"파일 해시 계산 완료: {file_hash} - {file_info.get('original_filename', '파일')}")
                return result, file_hash
            
            # 임시 파일 정리
            try:
                os.remove(file_path)
            except:
                pass
            
            return result, file_hash
        else:
            return f"[{field_name} - {file_info.get('original_filename', '파일')}]: 파일을 다운로드할 수 없습니다.", None
    except Exception as e:
        logger.error(f"파일 텍스트 추출 실패: {e}")
        return f"[{field_name} - {file_info.get('original_filename', '파일')}]: 텍스트 추출 중 오류 발생", None

@csrf_exempt
def ai_chat_cache_clear(request):
    """AI 채팅 캐시 삭제 엔드포인트"""
    if request.method != 'POST':
        return JsonResponse({'success': False, 'error': 'POST 요청만 지원합니다'})
    try:
        data = json.loads(request.body)
        row_id = data.get('row_id')
        if not row_id:
            return JsonResponse({'success': False, 'error': 'row_id가 필요합니다'})
        cache_key = f'ai_chat_row_{row_id}'
        if cache_key in request.session:
            del request.session[cache_key]
            request.session.modified = True
            return JsonResponse({'success': True, 'message': '캐시가 삭제되었습니다'})
        else:
            return JsonResponse({'success': True, 'message': '캐시가 존재하지 않습니다'})
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)})

@csrf_exempt
def file_cache_management(request):
    """파일 캐시 관리 API"""
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            action = data.get('action')
            
            if action == 'clear_all':
                success = clear_all_caches()
                return JsonResponse({
                    'success': success,
                    'message': '모든 캐시가 정리되었습니다.' if success else '캐시 정리 중 오류가 발생했습니다.'
                })
            elif action == 'get_stats':
                stats = get_cache_stats()
                return JsonResponse({
                    'success': True,
                    'stats': stats
                })
            else:
                return JsonResponse({'success': False, 'error': '지원하지 않는 액션입니다.'})
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})
    else:
        # GET 요청으로 캐시 통계 반환
        stats = get_cache_stats()
        return JsonResponse({
            'success': True,
            'stats': stats
        })

@csrf_exempt
def performance_monitoring(request):
    """성능 모니터링 API"""
    if request.method == 'GET':
        try:
            # 현재 캐시 상태
            cache_stats = get_cache_stats()
            
            # 성능 통계
            performance_stats = get_performance_stats()
            
            # 시스템 정보
            import psutil
            system_info = {
                'cpu_percent': psutil.cpu_percent(interval=1),
                'memory_percent': psutil.virtual_memory().percent,
                'disk_usage': psutil.disk_usage('/').percent,
                'memory_available_mb': psutil.virtual_memory().available / (1024 * 1024)
            }
            
            # 최적화 권장사항
            optimization_suggestions = []
            
            # 메모리 사용량 체크
            if system_info['memory_percent'] > 80:
                optimization_suggestions.append("메모리 사용량이 높습니다. 캐시를 정리하는 것을 권장합니다.")
            
            # API 응답 시간 체크
            if performance_stats.get('avg_api_time', 0) > 10:
                optimization_suggestions.append("OpenAI API 응답 시간이 느립니다. 네트워크 상태를 확인하세요.")
            
            # 캐시 히트율 체크
            cache_hit_rate = len(file_text_cache) / max(len(file_text_cache) + 1, 1)
            if cache_hit_rate < 0.3:
                optimization_suggestions.append("캐시 히트율이 낮습니다. 캐시 전략을 재검토하세요.")
            
            return JsonResponse({
                'success': True,
                'cache_stats': cache_stats,
                'performance_stats': performance_stats,
                'system_info': system_info,
                'optimization_suggestions': optimization_suggestions,
                'file_processing_pool_size': file_processing_pool._max_workers,
                'openai_pool_size': openai_pool._max_workers
            })
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})
    elif request.method == 'POST':
        try:
            data = json.loads(request.body)
            action = data.get('action')
            
            if action == 'optimize_memory':
                success = optimize_memory_usage()
                return JsonResponse({
                    'success': success,
                    'message': '메모리 최적화가 완료되었습니다.' if success else '메모리 최적화 중 오류가 발생했습니다.'
                })
            elif action == 'preload_files':
                success = preload_common_files()
                return JsonResponse({
                    'success': success,
                    'message': '파일 미리 로드가 완료되었습니다.' if success else '파일 미리 로드 중 오류가 발생했습니다.'
                })
            elif action == 'clear_metrics':
                performance_metrics.clear()
                performance_metrics.update({
                    'file_processing_times': [],
                    'openai_api_times': [],
                    'cache_hit_rates': []
                })
                return JsonResponse({
                    'success': True,
                    'message': '성능 메트릭이 초기화되었습니다.'
                })
            else:
                return JsonResponse({'success': False, 'error': '지원하지 않는 액션입니다.'})
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})
    else:
        return JsonResponse({'success': False, 'error': 'GET 또는 POST 요청만 지원합니다'})

def extract_text_from_file(file_path):
    """파일에서 텍스트 추출 (update_bizinfo.py 참고) - PDF 표 우선 추출, 텍스트 부족시 OCR 자동 fallback"""
    import os
    if not file_path or not os.path.exists(file_path):
        return ""
    
    def is_text_extracted_enough(file_path, extracted_text):
        file_size = os.path.getsize(file_path)
        text_length = len(extracted_text)
        if file_size == 0:
            return False
        ratio = text_length / file_size
        return ratio > 0.001  # 0.1% 이상이면 텍스트가 어느 정도 있다고 판단

    try:
        if file_path.endswith(".pdf"):
            import tempfile
            with pdfplumber.open(file_path) as pdf:
                all_tables = []
                for page in pdf.pages:
                    tables = page.extract_tables()
                    for table in tables:
                        all_tables.append(table)
                if all_tables:
                    # 표가 있으면 표만 추출
                    table_texts = []
                    for table in all_tables:
                        table_texts.append('\n'.join(['\t'.join([cell if cell is not None else '' for cell in row]) for row in table if row]))
                    result = '\n\n'.join(table_texts)
                else:
                    # 표가 없으면 기존 방식
                    text = ''
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + '\n'
                    result = text.strip()
            # 텍스트 부족시 OCR 자동 fallback
            if not is_text_extracted_enough(file_path, result):
                ocr_texts = []
                with pdfplumber.open(file_path) as pdf:
                    for i, page in enumerate(pdf.pages):
                        # 페이지를 이미지로 저장
                        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_img:
                            img = page.to_image(resolution=300)
                            img.save(tmp_img.name, format='PNG')
                            # OCR 수행
                            ocr_text = clova_ocr(tmp_img.name, 'jpg')
                            ocr_texts.append(f"[페이지 {i+1} OCR 결과]:\n{ocr_text}")
                            try:
                                os.remove(tmp_img.name)
                            except:
                                pass
                return f"[경고] 이 PDF는 텍스트가 거의 없는 이미지 기반 PDF로 판단되어 OCR로 텍스트를 추출했습니다.\n\n" + '\n\n'.join(ocr_texts)
            return result
        elif file_path.endswith((".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff", ".tif", ".webp")):
            return clova_ocr(file_path, "jpg")
        elif file_path.endswith((".hwp", ".hwpx")):
            pdf_path = convert_hwp_to_pdf(file_path)
            if os.path.exists(pdf_path):
                extracted_text = extract_text_from_file_optimized(pdf_path)
                if os.path.exists(pdf_path):
                    os.remove(pdf_path)
                return extracted_text
            else:
                return "HWP 파일 변환 실패"
        elif file_path.endswith((".docx", ".doc")):
            # DOCX 파일 처리
            try:
                from docx import Document
                
                # DOCX 파일 열기
                doc = Document(file_path)
                print(f"  DOCX 파일 열기 성공")
                
                # 모든 단락의 텍스트 추출
                paragraphs = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        paragraphs.append(paragraph.text.strip())
                        print(f"  단락 {paragraph.text.strip()[:100]}...")
                
                print(f"  추출된 단락 수: {len(paragraphs)}")
                
                # 모든 테이블의 텍스트 추출
                tables = []
                for table_idx, table in enumerate(doc.tables):
                    print(f"  테이블 {table_idx+1} 처리 중...")
                    table_text = []
                    for row_idx, row in enumerate(table.rows):
                        row_text = []
                        for cell_idx, cell in enumerate(row.cells):
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                                print(f"    셀 {row_idx+1}-{cell_idx+1}: {cell.text.strip()[:50]}...")
                        if row_text:
                            table_text.append(' | '.join(row_text))
                    if table_text:
                        tables.append('\n'.join(table_text))
                        print(f"  테이블 {table_idx+1} 텍스트 완성")
                
                print(f"  추출된 테이블 수: {len(tables)}")
                
                # 결과 조합
                result = []
                if paragraphs:
                    result.append('\n'.join(paragraphs))
                    print(f"  단락 텍스트 추가됨")
                if tables:
                    result.append('\n\n'.join(tables))
                    print(f"  테이블 텍스트 추가됨")
                
                final_result = '\n\n'.join(result) if result else "DOCX 파일에서 텍스트를 추출할 수 없습니다."
                
                print(f"  최종 결과 길이: {len(final_result)}")
                print(f"  최종 결과 미리보기: {final_result[:200]}...")
                
                return final_result
                
            except ImportError:
                logger.error("python-docx 라이브러리가 설치되지 않았습니다.")
                print(f"  오류: python-docx 라이브러리가 설치되지 않았습니다.")
                return "DOCX 파일 처리를 위해 python-docx 라이브러리가 필요합니다."
            except Exception as e:
                logger.error(f"DOCX 텍스트 추출 실패: {e}")
                print(f"  오류: DOCX 텍스트 추출 실패 - {e}")
                return f"DOCX 텍스트 추출 실패: {str(e)}"
        elif file_path.endswith((".txt", ".md", ".markdown", ".rst", ".adoc")):
            # 텍스트 파일과 마크다운 파일 처리
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                    return content.strip()
            except UnicodeDecodeError:
                # UTF-8로 읽기 실패시 다른 인코딩 시도
                try:
                    with open(file_path, 'r', encoding='cp949') as f:
                        content = f.read()
                        return content.strip()
                except UnicodeDecodeError:
                    try:
                        with open(file_path, 'r', encoding='euc-kr') as f:
                            content = f.read()
                            return content.strip()
                    except Exception as e:
                        logger.error(f"텍스트 파일 인코딩 처리 실패: {e}")
                        return f"텍스트 파일 읽기 실패: 인코딩 문제"
            except Exception as e:
                logger.error(f"텍스트 파일 읽기 실패: {e}")
                return f"텍스트 파일 읽기 실패: {str(e)}"
        elif file_path.endswith((".csv", ".tsv")):
            # CSV/TSV 파일 처리
            try:
                import csv
                delimiter = ',' if file_path.endswith('.csv') else '\t'
                with open(file_path, 'r', encoding='utf-8') as f:
                    reader = csv.reader(f, delimiter=delimiter)
                    rows = list(reader)
                    if rows:
                        # 헤더와 데이터를 텍스트로 변환
                        result = []
                        for i, row in enumerate(rows):
                            if i == 0:  # 헤더
                                result.append(f"헤더: {' | '.join(row)}")
                            else:  # 데이터
                                result.append(f"행 {i}: {' | '.join(row)}")
                        return '\n'.join(result)
                    else:
                        return "빈 CSV/TSV 파일"
            except Exception as e:
                logger.error(f"CSV/TSV 파일 읽기 실패: {e}")
                return f"CSV/TSV 파일 읽기 실패: {str(e)}"
        elif file_path.endswith((".json", ".xml", ".yaml", ".yml")):
            # 구조화된 데이터 파일 처리
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                    return f"구조화된 데이터 파일 내용:\n{content.strip()}"
            except Exception as e:
                logger.error(f"구조화된 데이터 파일 읽기 실패: {e}")
                return f"구조화된 데이터 파일 읽기 실패: {str(e)}"
        return ""
    except Exception as e:
        logger.error(f"텍스트 추출 실패: {e}")
        return ""

def is_text_pdf(file_path):
    """PDF가 텍스트 기반인지 확인"""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages[:2]:
                if page.extract_text():
                    return True
        return False
    except:
        return False

def clova_ocr(file_path, fmt):
    """Clova OCR을 사용한 텍스트 추출"""
    from config import NAVER_CLOVA_OCR_API_KEY, NAVER_CLOUD_CLOVA_OCR_API_URL
    
    print(f"Clova OCR 시작: {file_path}, 형식: {fmt}")
    
    request_json = {
        'images': [{'format': fmt, 'name': 'demo'}],
        'requestId': str(uuid.uuid4()),
        'version': 'V1',
        'timestamp': int(time.time() * 1000)
    }
    payload = {'message': json.dumps(request_json).encode('UTF-8')}
    
    try:
        with open(file_path, 'rb') as f:
            files = [('file', f)]
            headers = {'X-OCR-SECRET': NAVER_CLOVA_OCR_API_KEY}
            
            response = requests.post(NAVER_CLOUD_CLOVA_OCR_API_URL, headers=headers, data=payload, files=files)
            
            if response.status_code != 200:
                print(f"Clova OCR API 응답 오류: {response.status_code}, {response.text}")
                return ""
            
            response_data = response.json()
            if 'images' not in response_data or not response_data['images']:
                print(f"Clova OCR 응답에 이미지 데이터가 없음: {response_data}")
                return ""
            
            full_text = ""
            for field in response_data['images'][0].get('fields', []):
                full_text += field['inferText'] + " "
            
            extracted_text = full_text.strip()
            print(f"Clova OCR 완료: {len(extracted_text)} characters extracted")
            return extracted_text
            
    except FileNotFoundError:
        print(f"파일을 찾을 수 없음: {file_path}")
        return ""
    except requests.exceptions.RequestException as e:
        print(f"Clova OCR API 요청 실패: {e}")
        return ""
    except Exception as e:
        print(f"Clova OCR 처리 중 예상치 못한 오류: {e}")
        return ""

def check_libreoffice_status():
    """LibreOffice 설치 상태와 버전을 확인합니다."""
    try:
        result = subprocess.run([
            "libreoffice", "--version"
        ], stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=30)
        
        if result.returncode == 0:
            version = result.stdout.decode().strip()
            logger.info(f"✅ LibreOffice 설치 확인: {version}")
            return True
        else:
            logger.error(f"❌ LibreOffice 실행 실패: {result.stderr.decode()}")
            return False
    except Exception as e:
        logger.error(f"❌ LibreOffice 확인 실패: {e}")
        return False

def convert_hwp_to_pdf(hwp_path):
    """HWP를 PDF로 변환 (개선된 버전)"""
    output_dir = os.path.dirname(hwp_path)
    try:
        # 파일 크기 확인
        file_size = os.path.getsize(hwp_path)
        logger.info(f"📄 HWP 파일 크기: {file_size / (1024*1024):.2f} MB")
        
        # 파일 크기에 따른 timeout 조정
        if file_size > 50 * 1024 * 1024:  # 50MB 이상
            timeout = 1800  # 30분
            logger.info("⏰ 대용량 파일 감지, timeout을 30분으로 설정")
        elif file_size > 10 * 1024 * 1024:  # 10MB 이상
            timeout = 900   # 15분
            logger.info("⏰ 중간 크기 파일 감지, timeout을 15분으로 설정")
        else:
            timeout = 600   # 10분 (기본값)
            logger.info("⏰ 기본 timeout 10분 설정")
        
        # LibreOffice 프로세스 시작 전 메모리 상태 확인
        logger.info("🖥️ LibreOffice 변환 시작...")
        
        result = subprocess.run([
            "libreoffice",
            "--headless",
            "--convert-to", "pdf:writer_pdf_Export",
            hwp_path,
            "--outdir", output_dir
        ], stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=timeout)

        logger.info("🖥️ libreoffice stdout: " + result.stdout.decode())
        if result.stderr:
            logger.info("🖥️ libreoffice stderr: " + result.stderr.decode())

        basename = os.path.splitext(os.path.basename(hwp_path))[0] + ".pdf"
        converted_pdf = os.path.join(output_dir, basename)

        if os.path.exists(converted_pdf):
            pdf_size = os.path.getsize(converted_pdf)
            logger.info(f"✅ 변환 성공: {pdf_size / (1024*1024):.2f} MB")
            return converted_pdf
        else:
            logger.error(f"[❌ 변환 실패] {converted_pdf} 파일이 존재하지 않습니다.")
            return ""
            
    except subprocess.TimeoutExpired:
        logger.error(f"[⏰ Timeout 발생] {timeout}초 초과로 변환 실패")
        # LibreOffice 프로세스 강제 종료
        try:
            subprocess.run(["pkill", "-f", "libreoffice"], timeout=10)
            logger.info("🔄 LibreOffice 프로세스 강제 종료 완료")
        except:
            logger.warning("⚠️ LibreOffice 프로세스 종료 실패")
        return ""
    except Exception as e:
        logger.error(f"[예외 발생] HWP → PDF 변환 실패: {e}")
        return ""

def download_file_from_url(url):
    """URL에서 파일 다운로드"""
    try:
        import tempfile
        import requests
        
        # S3 URL인 경우 boto3 사용
        if 's3.ap-northeast-2.amazonaws.com' in url:
            return download_file_from_s3(url)
        
        # 일반 URL인 경우 requests 사용
        temp_dir = tempfile.gettempdir()
        file_name = url.split('/')[-1].split('?')[0]  # 쿼리 파라미터 제거
        temp_path = os.path.join(temp_dir, file_name)
        
        # 파일 다운로드
        response = requests.get(url, stream=True, timeout=30)
        response.raise_for_status()
        
        with open(temp_path, 'wb') as f:
            for chunk in response.iter_content(chunk_size=8192):
                f.write(chunk)
        
        return temp_path
    except Exception as e:
        logger.error(f"파일 다운로드 실패: {e}")
        return None

def download_file_from_s3(url):
    """S3에서 직접 파일 다운로드"""
    try:
        # S3 클라이언트 생성
        s3_client = boto3.client(
            's3',
            aws_access_key_id=AWS_S3_ACCESS_KEY,
            aws_secret_access_key=AWS_S3_SECRET_KEY,
            region_name=AWS_S3_REGION
        )
        
        # URL에서 S3 키 추출
        bucket_name = AWS_S3_BUCKET_NAME
        
        # URL에서 파일 경로 추출
        if '/media/' in url:
            s3_key = 'media/' + url.split('/media/')[-1].split('?')[0]
        elif '/note_files/' in url:
            s3_key = 'note_files/' + url.split('/note_files/')[-1].split('?')[0]
        else:
            # 다른 경로인 경우 전체 경로에서 추출
            s3_key = url.split(f'{bucket_name}/')[-1].split('?')[0]
        
        # 임시 파일 생성
        temp_dir = tempfile.gettempdir()
        file_name = s3_key.split('/')[-1]
        temp_path = os.path.join(temp_dir, file_name)
        
        # S3에서 파일 다운로드
        s3_client.download_file(bucket_name, s3_key, temp_path)
        
        return temp_path
    except Exception as e:
        logger.error(f"S3 파일 다운로드 실패: {e}")
        return None

def download_file_from_s3_key(s3_key):
    """S3 키를 사용하여 파일 다운로드"""
    try:
        # S3 클라이언트 생성
        s3_client = boto3.client(
            's3',
            aws_access_key_id=AWS_S3_ACCESS_KEY,
            aws_secret_access_key=AWS_S3_SECRET_KEY,
            region_name=AWS_S3_REGION
        )
        
        # 임시 파일 생성
        temp_dir = tempfile.gettempdir()
        file_name = s3_key.split('/')[-1]
        temp_path = os.path.join(temp_dir, file_name)
        
        # S3에서 파일 다운로드
        s3_client.download_file(AWS_S3_BUCKET_NAME, s3_key, temp_path)
        
        return temp_path
    except Exception as e:
        logger.error(f"S3 키를 사용한 파일 다운로드 실패: {e}")
        return None

def process_files_parallel(file_infos, field_name):
    """병렬로 여러 파일 처리"""
    try:
        # 스레드 풀을 사용하여 병렬 처리
        futures = []
        for file_info in file_infos:
            future = file_processing_pool.submit(process_file_async, file_info, field_name)
            futures.append(future)
        
        # 결과 수집
        results = []
        for future in concurrent.futures.as_completed(futures, timeout=60):
            try:
                file_text, file_hash = future.result()
                if file_text:
                    results.append(file_text)
            except Exception as e:
                logger.error(f"병렬 파일 처리 실패: {e}")
        
        return results
    except Exception as e:
        logger.error(f"병렬 파일 처리 실패: {e}")
        return []

def cleanup_temp_files():
    """임시 파일 정리 (개선된 버전)"""
    try:
        temp_dir = tempfile.gettempdir()
        current_time = time.time()
        cleaned_count = 0
        
        # 임시 파일들 정리
        for filename in os.listdir(temp_dir):
            if filename.startswith(('s3_', 'url_')):
                file_path = os.path.join(temp_dir, filename)
                try:
                    file_age = current_time - os.path.getmtime(file_path)
                    if file_age > 3600:  # 1시간
                        os.remove(file_path)
                        cleaned_count += 1
                except:
                    pass
        
        # 캐시 디렉토리 정리
        if CACHE_DIR.exists():
            for cache_file in CACHE_DIR.glob("*.pkl"):
                try:
                    file_age = current_time - cache_file.stat().st_mtime
                    if file_age > 86400:  # 24시간
                        cache_file.unlink()
                        cleaned_count += 1
                except:
                    pass
        
        if cleaned_count > 0:
            logger.info(f"임시 파일 정리 완료: {cleaned_count}개 파일 삭제")
            
    except Exception as e:
        logger.error(f"임시 파일 정리 실패: {e}")

def monitor_performance(func):
    """성능 모니터링 데코레이터"""
    def wrapper(*args, **kwargs):
        start_time = time.time()
        start_memory = os.getpid()  # 간단한 메모리 추적
        
        try:
            result = func(*args, **kwargs)
            end_time = time.time()
            execution_time = end_time - start_time
            
            # 성능 로깅
            if execution_time > 1.0:  # 1초 이상 걸리는 작업만 로깅
                logger.info(f"성능 모니터링 - {func.__name__}: {execution_time:.2f}초")
            
            return result
        except Exception as e:
            end_time = time.time()
            execution_time = end_time - start_time
            logger.error(f"성능 모니터링 - {func.__name__} 실패: {execution_time:.2f}초, 오류: {e}")
            raise
    
    return wrapper

@monitor_performance
def extract_text_from_file_optimized(file_path):
    """최적화된 파일 텍스트 추출 (성능 모니터링 포함)"""
    if not file_path or not os.path.exists(file_path):
        return ""
    
    # 파일 크기 체크
    file_size = os.path.getsize(file_path)
    
    # 매우 큰 파일은 처리 제한
    if file_size > 50 * 1024 * 1024:  # 50MB
        return f"파일이 너무 커서 텍스트 추출을 건너뜁니다. (크기: {file_size / (1024*1024):.1f}MB)"
    
    # 파일 해시 계산
    file_hash = calculate_file_hash_fast(file_path)
    
    # 캐시 확인
    cached_text = get_cached_file_text(file_hash, file_path)
    if cached_text:
        return cached_text
    
    try:
        if file_path.endswith(".pdf"):
            return extract_pdf_text_optimized(file_path, file_hash)
        elif file_path.endswith((".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff", ".tif", ".webp")):
            return extract_image_text_optimized(file_path, file_hash)
        elif file_path.endswith(".hwp"):
            return extract_hwp_text_optimized(file_path, file_hash)
        elif file_path.endswith((".txt", ".md", ".markdown", ".rst", ".adoc")):
            return extract_text_file_optimized(file_path, file_hash)
        elif file_path.endswith((".csv", ".tsv")):
            return extract_csv_text_optimized(file_path, file_hash)
        elif file_path.endswith((".json", ".xml", ".yaml", ".yml")):
            return extract_structured_text_optimized(file_path, file_hash)
        return ""
    except Exception as e:
        logger.error(f"최적화된 텍스트 추출 실패: {e}")
        return ""

def get_cache_stats():
    """캐시 통계 정보 반환"""
    try:
        memory_cache_size = len(file_text_cache)
        disk_cache_files = len(list(CACHE_DIR.glob("*.pkl"))) if CACHE_DIR.exists() else 0
        dropdown_cache_size = len(dropdown_cache)  # 드롭다운 캐시 크기 추가
        
        # 디스크 캐시 크기 계산
        disk_cache_size = 0
        if CACHE_DIR.exists():
            for cache_file in CACHE_DIR.glob("*.pkl"):
                disk_cache_size += cache_file.stat().st_size
        
        return {
            'memory_cache_size': memory_cache_size,
            'disk_cache_files': disk_cache_files,
            'disk_cache_size_mb': disk_cache_size / (1024 * 1024),
            'dropdown_cache_size': dropdown_cache_size,  # 드롭다운 캐시 크기 추가
            'max_cache_size': MAX_CACHE_SIZE
        }
    except Exception as e:
        logger.error(f"캐시 통계 조회 실패: {e}")
        return {}

def clear_all_caches():
    """모든 캐시 정리"""
    try:
        # 메모리 캐시 정리
        file_text_cache.clear()
        file_hash_cache.clear()
        dropdown_cache.clear()  # 드롭다운 캐시도 정리
        
        # 디스크 캐시 정리
        if CACHE_DIR.exists():
            for cache_file in CACHE_DIR.glob("*.pkl"):
                try:
                    cache_file.unlink()
                except:
                    pass
        
        # 임시 파일 정리
        cleanup_temp_files()
        
        logger.info("모든 캐시가 정리되었습니다.")
        return True
    except Exception as e:
        logger.error(f"캐시 정리 실패: {e}")
        return False

def cleanup_related_caches(deleted_files, added_files, modified_files):
    """파일 변경사항에 관련된 캐시 정리"""
    try:
        cleaned_count = 0
        
        # 삭제된 파일들의 캐시 정리
        for deleted_file in deleted_files:
            file_info = deleted_file.get('fileInfo', {})
            if file_info:
                # 파일 해시 기반 캐시 정리
                file_hash = file_info.get('file_hash')
                if file_hash:
                    # 메모리 캐시에서 제거
                    if file_hash in file_text_cache:
                        del file_text_cache[file_hash]
                        cleaned_count += 1
                    
                    # 디스크 캐시에서 제거
                    cache_file = CACHE_DIR / f"{file_hash}.pkl"
                    if cache_file.exists():
                        try:
                            cache_file.unlink()
                            cleaned_count += 1
                        except Exception as e:
                            logger.error(f"디스크 캐시 제거 실패: {e}")
                
                # s3_key 기반 캐시 정리
                s3_key = file_info.get('s3_key')
                if s3_key:
                    cache_key = f"s3_download_{hashlib.md5(s3_key.encode()).hexdigest()}"
                    if cache_key in file_hash_cache:
                        cached_path = file_hash_cache[cache_key]
                        if os.path.exists(cached_path):
                            try:
                                os.remove(cached_path)
                                cleaned_count += 1
                            except Exception as e:
                                logger.error(f"S3 캐시 파일 제거 실패: {e}")
                        del file_hash_cache[cache_key]
                        cleaned_count += 1
        
        # 수정된 파일들의 캐시 정리 (기존 캐시 무효화)
        for modified_file in modified_files:
            file_info = modified_file.get('fileInfo', {})
            if file_info:
                # 파일 해시 기반 캐시 정리
                file_hash = file_info.get('file_hash')
                if file_hash:
                    # 메모리 캐시에서 제거 (새로 생성될 예정)
                    if file_hash in file_text_cache:
                        del file_text_cache[file_hash]
                        cleaned_count += 1
                    
                    # 디스크 캐시에서 제거
                    cache_file = CACHE_DIR / f"{file_hash}.pkl"
                    if cache_file.exists():
                        try:
                            cache_file.unlink()
                            cleaned_count += 1
                        except Exception as e:
                            logger.error(f"디스크 캐시 제거 실패: {e}")
        
        if cleaned_count > 0:
            logger.info(f"관련 캐시 정리 완료: {cleaned_count}개 항목 제거")
            print(f"관련 캐시 정리 완료: {cleaned_count}개 항목 제거")
        
    except Exception as e:
        logger.error(f"관련 캐시 정리 실패: {e}")
        print(f"관련 캐시 정리 실패: {e}")

# 주기적으로 임시 파일 정리 (서버 시작 시)
cleanup_temp_files()

# 성능 모니터링을 위한 주기적 캐시 정리 (24시간마다)
def schedule_cache_cleanup():
    """주기적 캐시 정리 스케줄러"""
    import threading
    import time
    
    def cleanup_worker():
        while True:
            try:
                time.sleep(86400)  # 24시간 대기
                cleanup_temp_files()
                
                # 캐시 크기 체크 및 정리
                stats = get_cache_stats()
                if stats.get('memory_cache_size', 0) > MAX_CACHE_SIZE * 0.8:
                    # 메모리 캐시가 80% 이상 차면 오래된 항목들 정리
                    oldest_keys = sorted(file_text_cache.keys())[:MAX_CACHE_SIZE // 4]
                    for key in oldest_keys:
                        del file_text_cache[key]
                    logger.info(f"메모리 캐시 정리: {len(oldest_keys)}개 항목 제거")
                
            except Exception as e:
                logger.error(f"주기적 캐시 정리 실패: {e}")
    
    # 백그라운드 스레드로 실행
    cleanup_thread = threading.Thread(target=cleanup_worker, daemon=True)
    cleanup_thread.start()

# 캐시 정리 스케줄러 시작
schedule_cache_cleanup()

def get_file_size_category(file_size):
    """파일 크기에 따른 처리 카테고리 결정"""
    if file_size <= FILE_SIZE_THRESHOLDS['small']:
        return 'small'
    elif file_size <= FILE_SIZE_THRESHOLDS['medium']:
        return 'medium'
    else:
        return 'large'

def optimize_file_processing_strategy(files):
    """파일 크기별 최적화된 처리 전략"""
    small_files = []
    medium_files = []
    large_files = []
    
    for file_info in files:
        file_size = file_info.get('file_size', 0)
        category = get_file_size_category(file_size)
        
        if category == 'small':
            small_files.append(file_info)
        elif category == 'medium':
            medium_files.append(file_info)
        else:
            large_files.append(file_info)
    
    return {
        'small': small_files,
        'medium': medium_files,
        'large': large_files
    }

def process_files_with_priority(file_groups, field_name):
    """우선순위 기반 파일 처리"""
    results = []
    
    # 1. 작은 파일들을 먼저 병렬 처리 (빠른 응답)
    if file_groups['small']:
        print(f"작은 파일 병렬 처리: {len(file_groups['small'])}개")
        small_results = process_files_parallel(file_groups['small'], field_name)
        results.extend(small_results)
    
    # 2. 중간 크기 파일들을 병렬 처리
    if file_groups['medium']:
        print(f"중간 파일 병렬 처리: {len(file_groups['medium'])}개")
        medium_results = process_files_parallel(file_groups['medium'], field_name)
        results.extend(medium_results)
    
    # 3. 큰 파일들을 순차 처리 (메모리 효율성)
    if file_groups['large']:
        print(f"큰 파일 순차 처리: {len(file_groups['large'])}개")
        for file_info in file_groups['large']:
            try:
                file_text, file_hash = extract_file_text(field_name, file_info)
                if file_text:
                    results.append(file_text)
            except Exception as e:
                logger.error(f"큰 파일 처리 실패: {e}")
    
    return results

def call_openai_api_optimized(messages, retry_count=0):
    """최적화된 OpenAI API 호출 (재시도 로직 포함)"""
    start_time = time.time()
    
    try:
        headers = {
            'Authorization': f'Bearer {OPEN_AI_API_KEY}',
            'Content-Type': 'application/json'
        }
        
        payload = {
            'model': 'gpt-4.1-mini',  # 더 빠른 모델 사용
            'messages': messages,
            'max_tokens': 2000,  # 1500 -> 2000으로 증가
            'temperature': 0.3,  # 0.7 -> 0.3으로 감소 (더 정확한 응답)
            'stream': False  # 스트리밍 비활성화로 응답 속도 향상
        }
        
        response = requests.post(
            'https://api.openai.com/v1/chat/completions',
            headers=headers,
            json=payload,
            timeout=OPENAI_API_CONFIG['timeout']
        )
        
        if response.status_code == 200:
            result = response.json()
            ai_response = result['choices'][0]['message']['content'].strip()
            
            # 성능 메트릭 기록
            execution_time = time.time() - start_time
            performance_metrics['openai_api_times'].append(execution_time)
            
            return {'success': True, 'response': ai_response, 'time': execution_time}
        else:
            error_msg = f"OpenAI API 오류: {response.status_code}"
            try:
                error_data = response.json()
                if 'error' in error_data:
                    error_msg = f"OpenAI API 오류: {error_data['error'].get('message', '알 수 없는 오류')}"
            except:
                pass
            
            # 재시도 로직
            if retry_count < OPENAI_API_CONFIG['max_retries']:
                time.sleep(OPENAI_API_CONFIG['retry_delay'] * (retry_count + 1))
                return call_openai_api_optimized(messages, retry_count + 1)
            
            return {'success': False, 'error': error_msg}
            
    except requests.exceptions.Timeout:
        if retry_count < OPENAI_API_CONFIG['max_retries']:
            time.sleep(OPENAI_API_CONFIG['retry_delay'] * (retry_count + 1))
            return call_openai_api_optimized(messages, retry_count + 1)
        return {'success': False, 'error': '요청 시간이 초과되었습니다.'}
    except Exception as e:
        return {'success': False, 'error': f'API 호출 실패: {str(e)}'}

def optimize_context_for_openai(row_data, file_texts, max_context_length=12000):
    """OpenAI API용 컨텍스트 최적화 (정확성 향상)"""
    context_parts = []
    
    # 행 데이터 최적화 (더 많은 정보 포함)
    if row_data:
        row_summary = []
        for key, value in row_data.items():
            # 길이 제한을 늘려서 더 많은 정보 포함
            if value and len(str(value)) < 500:  # 200 -> 500으로 증가
                row_summary.append(f"{key}: {value}")
        
        if row_summary:
            context_parts.append("행 데이터:\n" + "\n".join(row_summary[:30]))  # 20 -> 30으로 증가
    
    # 파일 텍스트 최적화 (더 많은 텍스트 포함)
    if file_texts:
        file_summaries = []
        total_length = 0
        
        for file_text in file_texts:
            # 파일 텍스트 길이 제한 증가
            if len(file_text) > 4000:  # 2000 -> 4000으로 증가
                file_text = file_text[:4000] + "...[파일이 너무 커서 일부만 표시]"
            
            if total_length + len(file_text) < max_context_length:
                file_summaries.append(file_text)
                total_length += len(file_text)
            else:
                # 컨텍스트가 가득 찰 때는 중요한 파일부터 우선순위 부여
                break
        
        if file_summaries:
            context_parts.append("첨부 파일:\n" + "\n".join(file_summaries))
    
    return "\n\n".join(context_parts)

def batch_process_files(files, batch_size=OPENAI_API_CONFIG['batch_size']):
    """파일들을 배치로 나누어 처리"""
    batches = []
    for i in range(0, len(files), batch_size):
        batches.append(files[i:i + batch_size])
    return batches

def get_performance_stats():
    """성능 통계 정보 반환"""
    try:
        if not performance_metrics['openai_api_times']:
            return {}
        
        api_times = performance_metrics['openai_api_times']
        return {
            'avg_api_time': sum(api_times) / len(api_times),
            'min_api_time': min(api_times),
            'max_api_time': max(api_times),
            'total_api_calls': len(api_times),
            'recent_api_calls': len(api_times[-10:]) if len(api_times) >= 10 else len(api_times)
        }
    except Exception as e:
        logger.error(f"성능 통계 조회 실패: {e}")
        return {}

def optimize_memory_usage():
    """메모리 사용량 최적화"""
    try:
        # 메모리 캐시 크기 제한
        if len(file_text_cache) > MAX_CACHE_SIZE:
            # 가장 오래된 항목들 제거
            oldest_keys = sorted(file_text_cache.keys())[:MAX_CACHE_SIZE // 4]
            for key in oldest_keys:
                del file_text_cache[key]
            logger.info(f"메모리 캐시 정리: {len(oldest_keys)}개 항목 제거")
        
        # 성능 메트릭 크기 제한
        for metric_name, metric_data in performance_metrics.items():
            if len(metric_data) > 1000:  # 최대 1000개 항목만 유지
                performance_metrics[metric_name] = metric_data[-1000:]
        
        return True
    except Exception as e:
        logger.error(f"메모리 최적화 실패: {e}")
        return False

def preload_common_files():
    """자주 사용되는 파일들을 미리 캐시에 로드"""
    try:
        # 자주 사용되는 파일 패턴들
        common_patterns = [
            '*.pdf',
            '*.docx',
            '*.txt'
        ]
        
        # 임시 디렉토리에서 자주 사용되는 파일들 찾기
        temp_dir = tempfile.gettempdir()
        for pattern in common_patterns:
            import glob
            files = glob.glob(os.path.join(temp_dir, pattern))
            for file_path in files[:5]:  # 최대 5개씩만
                try:
                    if os.path.exists(file_path):
                        file_hash = calculate_file_hash_fast(file_path)
                        if file_hash and file_hash not in file_text_cache:
                            # 백그라운드에서 미리 로드
                            file_processing_pool.submit(extract_text_from_file_optimized, file_path)
                except Exception as e:
                    logger.error(f"파일 미리 로드 실패: {e}")
        
        logger.info("자주 사용되는 파일 미리 로드 완료")
        return True
    except Exception as e:
        logger.error(f"파일 미리 로드 실패: {e}")
        return False

def get_dropdown_name(attribute, value):
    """드롭다운 값에서 실제 이름을 조회하는 헬퍼 함수"""
    cache_key = f"{attribute.id}_{value}"
    
    # 캐시 확인
    if cache_key in dropdown_cache:
        return dropdown_cache[cache_key]
    
    try:
        # ID로 직접 조회 (DropdownAttribute에는 value 필드가 없고 option 필드가 있음)
        dropdown_attr = DropdownAttribute.objects.get(
            attribute=attribute,
            id=value
        )
        result = dropdown_attr.option  # value 대신 option 필드 사용
        dropdown_cache[cache_key] = result
        return result
    except (DropdownAttribute.DoesNotExist, ValueError):
        # 조회 실패시 원본 값 사용
        result = f"{value} (코드의 의미는 데이터에 없음)"
        dropdown_cache[cache_key] = result
        return result