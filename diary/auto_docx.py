import json
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods
from .models import User, Row, AttributeValue
import json
from docx import Document
from docx.table import _Cell
from docx.text.paragraph import Paragraph
import re
from typing import Dict, List, Tuple
import requests
from config import OPEN_AI_API_KEY
import os
import tempfile
import subprocess
from PIL import Image
import warnings
import boto3
from botocore.exceptions import ClientError
import hashlib
import time
import uuid
import pdfplumber
import mmap
from io import BytesIO
import mimetypes

warnings.filterwarnings("ignore", category=UserWarning)

@csrf_exempt
@require_http_methods(["GET", "POST"])
def auto_docx_recommend(request):
    """OpenAI를 통한 사업 개요 추천 API"""
    print("=== auto_docx_recommend 함수 호출됨 ===")
    
    try:
        # 세션에서 사용자 ID 가져오기
        user_id = request.session.get('diary_member_id')
        if not user_id:
            return JsonResponse({
                'success': False,
                'error': '로그인이 필요합니다.'
            })
        
        # 사용자 정보 조회
        try:
            user = User.objects.get(id=user_id)
        except User.DoesNotExist:
            return JsonResponse({
                'success': False,
                'error': '사용자를 찾을 수 없습니다.'
            })
        
        print(f"사용자: {user.name} (ID: {user_id})")
        
        if request.method == "POST":
            try:
                print(f"요청 본문: {request.body}")
                data = json.loads(request.body)
                print(f"파싱된 데이터: {data}")
                
                row_id = data.get('row_id')
                service_product = data.get('service_product', '').strip()
                
                print(f"추출된 값들:")
                print(f"  - row_id: {row_id} (타입: {type(row_id)})")
                print(f"  - service_product: '{service_product}' (타입: {type(service_product)})")
                
                if not row_id:
                    return JsonResponse({
                        'success': False,
                        'error': 'row_id가 필요합니다.'
                    })
                
                if not service_product:
                    return JsonResponse({
                        'success': False,
                        'error': f'주 서비스·생산품목이 필요합니다. (받은 값: "{service_product}")'
                    })
                
                print(f"요청된 row_id: {row_id}")
                print(f"주 서비스·생산품목: {service_product}")
                
                # 특정 행 조회
                try:
                    row = Row.objects.filter(id=row_id, user=user).select_related('user').prefetch_related(
                        'values__attribute__attributeType',
                        'values__attribute__dropdown_attributes'
                    ).first()
                    
                    if not row:
                        return JsonResponse({
                            'success': False,
                            'error': '해당 행을 찾을 수 없습니다.'
                        })
                    
                    print(f"행 데이터 조회 성공: {row.id}")
                    
                    # 행의 모든 속성값들을 딕셔너리로 변환
                    row_data = {}
                    for attr_value in row.values.all():
                        if attr_value.attribute:
                            attr_name = attr_value.attribute.name
                            attr_value_text = attr_value.value
                            row_data[attr_name] = attr_value_text
                    
                    print(f"행 데이터: {row_data}")
                    
                    # 파일 텍스트 추출
                    file_texts = get_row_files_text(row)
                    print(f"추출된 파일 텍스트 개수: {len(file_texts)}")
                    
                    # OpenAI API를 통해 사업 개요 추천 생성
                    recommendation_data = generate_business_overview_with_openai(
                        row_data, service_product, row, file_texts
                    )
                    
                    if not recommendation_data:
                        return JsonResponse({
                            'success': False,
                            'error': '사업 개요 추천 생성에 실패했습니다.'
                        })
                    
                    print(f"OpenAI 추천 응답: {recommendation_data}")
                    
                    return JsonResponse({
                        'success': True,
                        'business_overview': recommendation_data.get('business_overview', '')
                    })
                    
                except Row.DoesNotExist:
                    return JsonResponse({
                        'success': False,
                        'error': '해당 행을 찾을 수 없습니다.'
                    })
                    
            except json.JSONDecodeError:
                return JsonResponse({
                    'success': False,
                    'error': '잘못된 JSON 형식입니다.'
                })
        else:
            return JsonResponse({
                'success': False,
                'error': 'POST 요청만 허용됩니다.'
            })
        
    except Exception as e:
        print(f"오류 발생: {str(e)}")
        return JsonResponse({
            'success': False,
            'error': str(e)
        })

@csrf_exempt
@require_http_methods(["GET", "POST"])
def auto_docx(request):
    print("=== auto_docx 함수 호출됨 ===")
    
    try:
        # 세션에서 사용자 ID 가져오기
        user_id = request.session.get('diary_member_id')
        if not user_id:
            return JsonResponse({
                'success': False,
                'error': '로그인이 필요합니다.'
            })
        
        # 사용자 정보 조회
        try:
            user = User.objects.get(id=user_id)
        except User.DoesNotExist:
            return JsonResponse({
                'success': False,
                'error': '사용자를 찾을 수 없습니다.'
            })
        
        print(f"사용자: {user.name} (ID: {user_id})")
        
        if request.method == "POST":
            # POST 요청: 특정 행의 DOCX 생성
            try:
                data = json.loads(request.body)
                row_id = data.get('row_id')
                service_product = data.get('service_product', '').strip()
                business_overview = data.get('business_overview', '').strip()
                
                if not row_id:
                    return JsonResponse({
                        'success': False,
                        'error': 'row_id가 필요합니다.'
                    })
                
                if not service_product:
                    return JsonResponse({
                        'success': False,
                        'error': '주 서비스·생산품목이 필요합니다.'
                    })
                
                print(f"요청된 row_id: {row_id}")
                print(f"주 서비스·생산품목: {service_product}")
                print(f"사업 개요: {business_overview}")
                
                # 특정 행 조회
                try:
                    row = Row.objects.filter(id=row_id, user=user).select_related('user').prefetch_related(
                        'values__attribute__attributeType',
                        'values__attribute__dropdown_attributes'
                    ).first()
                    
                    if not row:
                        return JsonResponse({
                            'success': False,
                            'error': '해당 행을 찾을 수 없습니다.'
                        })
                    
                    print(f"행 데이터 조회 성공: {row.id}")
                    
                    # 행의 모든 속성값들을 딕셔너리로 변환
                    row_data = {}
                    for attr_value in row.values.all():
                        if attr_value.attribute:
                            attr_name = attr_value.attribute.name
                            attr_value_text = attr_value.value
                            row_data[attr_name] = attr_value_text
                    
                    print(f"행 데이터: {row_data}")
                    
                    # 기업 정보에서 매출액 추출
                    revenue = row_data.get("매출액", row_data.get("매출액(백만원)", 0))

                    # 파일 텍스트 추출
                    file_texts = get_row_files_text(row)
                    print(f"추출된 파일 텍스트 개수: {len(file_texts)}")

                    # 사용자 입력값을 포함한 사업계획서 데이터 생성
                    business_plan_data = {
                        "주 서비스·생산품목": service_product,
                        "매출액(백만원)": revenue,
                        "사업개요": business_overview or "정보 부족",
                        "주 사용 플랫폼": row_data.get("주 사용 플랫폼", "정보 부족"),
                        "점포 보유현황": row_data.get("점포 보유현황", "정보 부족"),
                        "대표자_경력": row_data.get("대표자_경력", []),
                        "실제경영자_관계": row_data.get("실제경영자_관계", "본인"),
                        "실제경영자_경력": row_data.get("실제경영자_경력", []),
                        "향후 사업계획(자금용도 포함)": row_data.get("향후 사업계획", "정보 부족"),
                        "자금소요_운전자금": row_data.get("자금소요_운전자금", 30),
                        "자금조달_본건차입금": row_data.get("자금조달_본건차입금", 20),
                        "자금조달_자체자금": row_data.get("자금조달_자체자금", 10),
                        "자금조달_기타": row_data.get("자금조달_기타", 0),
                        "파일_텍스트": file_texts  # 파일 텍스트 추가
                    }
                    
                    print(f"사업계획서 데이터: {business_plan_data}")
                    
                    # DOCX 파일 생성
                    doc = Document("diary/사업계획서_양식.docx")
                    print(f"템플릿 파일 로드 완료: diary/사업계획서_양식.docx")
                    
                    # 사업계획서 데이터로 DOCX 채우기
                    try:
                        fill_business_plan_docx(doc, business_plan_data)
                        print("DOCX 채우기 완료")
                    except Exception as e:
                        print(f"DOCX 채우기 중 오류: {e}")
                        # 오류가 발생해도 기본 정보는 채워보기
                        try:
                            # 기본 정보만 채우기
                            basic_map = {
                                "주 서비스·생산품목": business_plan_data.get("주 서비스·생산품목", "정보 부족"),
                                "매출액": f'{business_plan_data.get("매출액(백만원)", 0)} 백만원',
                                "주 사용 플랫폼": business_plan_data.get("주 사용 플랫폼", "정보 부족"),
                            }
                            find_and_fill_simple_labels(doc, basic_map)
                            print("기본 정보 채우기 완료")
                        except Exception as basic_error:
                            print(f"기본 정보 채우기도 실패: {basic_error}")
                    
                    # 파일명 생성 (한글 안전 처리)
                    company_name = get_company_name_from_row(row)
                    print(f"회사명 추출: {company_name}")
                    
                    # 파일명에서 특수문자 제거 및 안전한 파일명 생성
                    safe_company_name = re.sub(r'[<>:"/\\|?*]', '_', str(company_name))
                    safe_company_name = safe_company_name.strip()
                    
                    if not safe_company_name:
                        safe_company_name = f"회사_{row.id}"
                    
                    output_filename = f"사업계획서_{safe_company_name}.docx"
                    
                    # 파일명이 너무 길면 자르기
                    if len(output_filename) > 100:
                        output_filename = f"사업계획서_{safe_company_name[:50]}.docx"
                    
                    print(f"생성될 파일명: {output_filename}")
                    
                    # 임시 파일로 저장 (다운로드용)
                    import tempfile
                    import os
                    
                    temp_dir = tempfile.gettempdir()
                    temp_file_path = os.path.join(temp_dir, output_filename)
                    print(f"임시 파일 경로: {temp_file_path}")
                    
                    try:
                        doc.save(temp_file_path)
                        print(f"임시 파일 저장 완료: {temp_file_path}")
                    except Exception as save_error:
                        print(f"임시 파일 저장 실패: {save_error}")
                        return JsonResponse({
                            'success': False,
                            'error': f'파일 저장에 실패했습니다: {str(save_error)}'
                        })
                    
                    # 파일이 실제로 생성되었는지 확인
                    if not os.path.exists(temp_file_path):
                        print(f"임시 파일이 생성되지 않음: {temp_file_path}")
                        return JsonResponse({
                            'success': False,
                            'error': '파일 생성에 실패했습니다.'
                        })
                    
                    file_size = os.path.getsize(temp_file_path)
                    print(f"생성된 파일 크기: {file_size} bytes")
                    
                    if file_size == 0:
                        print("생성된 파일이 비어있음")
                        return JsonResponse({
                            'success': False,
                            'error': '생성된 파일이 비어있습니다.'
                        })
                    
                    # 파일을 읽어서 HTTP 응답으로 전송
                    from django.http import FileResponse
                    import mimetypes
                     
                    try:
                        # 파일이 실제로 존재하는지 확인
                        if not os.path.exists(temp_file_path):
                            return JsonResponse({
                                'success': False,
                                'error': '생성된 파일을 찾을 수 없습니다.'
                            })
                        
                        # 파일 크기 확인
                        file_size = os.path.getsize(temp_file_path)
                        if file_size == 0:
                            return JsonResponse({
                                'success': False,
                                'error': '생성된 파일이 비어있습니다.'
                            })
                        
                        print(f"파일 다운로드 준비: {output_filename} (크기: {file_size} bytes)")
                        
                        # MIME 타입 설정
                        mime_type, _ = mimetypes.guess_type(output_filename)
                        if mime_type is None:
                            mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                        
                        # 파일 응답 생성 (더 안전한 방식)
                        try:
                            # 파일을 바이너리로 읽어서 메모리에 저장
                            with open(temp_file_path, 'rb') as file_handle:
                                file_content = file_handle.read()
                            
                            print(f"파일 읽기 완료: {len(file_content)} bytes")
                            
                            # BytesIO를 사용하여 메모리에서 파일 응답 생성
                            from io import BytesIO
                            file_stream = BytesIO(file_content)
                            file_stream.seek(0)  # 스트림 포인터를 처음으로
                            
                            response = FileResponse(
                                file_stream,
                                content_type=mime_type
                            )
                            
                            # 다운로드 강제 설정 (더 호환성 있는 방식)
                            safe_filename = output_filename.encode('utf-8').decode('latin-1')
                            response['Content-Disposition'] = f'attachment; filename="{safe_filename}"; filename*=UTF-8\'\'{output_filename}'
                            response['Content-Length'] = str(file_size)
                            response['Cache-Control'] = 'no-cache, no-store, must-revalidate'
                            response['Pragma'] = 'no-cache'
                            response['Expires'] = '0'
                            
                            print(f"파일 응답 생성 완료: {output_filename}")
                            print(f"=== 다운로드 준비 완료 ===")
                            print(f"파일명: {output_filename}")
                            print(f"파일 크기: {file_size} bytes")
                            print(f"MIME 타입: {mime_type}")
                            print(f"임시 경로: {temp_file_path}")
                            print(f"Content-Disposition: {response['Content-Disposition']}")
                            print(f"Content-Length: {response['Content-Length']}")
                            print(f"사용자에게 다운로드 시작...")
                            
                            # 임시 파일 정리 (응답 후)
                            def cleanup_temp_file():
                                try:
                                    if os.path.exists(temp_file_path):
                                        os.remove(temp_file_path)
                                        print(f"임시 파일 정리 완료: {temp_file_path}")
                                except Exception as cleanup_error:
                                    print(f"임시 파일 정리 실패: {cleanup_error}")
                            
                            # 응답 객체에 정리 함수 추가
                            response._cleanup = cleanup_temp_file
                            
                            return response
                                
                        except Exception as file_error:
                            print(f"파일 응답 생성 중 오류: {file_error}")
                            # 임시 파일 정리
                            try:
                                if os.path.exists(temp_file_path):
                                    os.remove(temp_file_path)
                            except:
                                pass
                            
                            return JsonResponse({
                                'success': False,
                                'error': f'파일 다운로드 준비 중 오류가 발생했습니다: {str(file_error)}'
                            })
                            
                    except Exception as download_error:
                        print(f"다운로드 준비 중 오류: {download_error}")
                        # 임시 파일 정리
                        try:
                            if os.path.exists(temp_file_path):
                                os.remove(temp_file_path)
                        except:
                            pass
                        
                        return JsonResponse({
                            'success': False,
                            'error': f'다운로드 준비 중 오류가 발생했습니다: {str(download_error)}'
                        })
                    
                except Row.DoesNotExist:
                    return JsonResponse({
                        'success': False,
                        'error': '해당 행을 찾을 수 없습니다.'
                    })
                    
            except json.JSONDecodeError:
                return JsonResponse({
                    'success': False,
                    'error': '잘못된 JSON 형식입니다.'
                })
                
        else:
            # GET 요청: 사용자의 모든 행 데이터 조회 (기존 기능 유지)
            rows = Row.objects.filter(user=user).select_related('user').prefetch_related(
                'values__attribute__attributeType',
                'values__attribute__dropdown_attributes'
            )
            
            print(f"총 {rows.count()}개의 행을 찾았습니다.")
            
            rows_data = []
            for row in rows[:10]:
                row_info = {
                    'id': row.id,
                    'created_at': row.created_at.isoformat() if row.created_at else None,
                    'attributes': {}
                }
                
                for attr_value in row.values.all():
                    if attr_value.attribute:
                        attr_name = attr_value.attribute.name
                        attr_value_text = attr_value.value
                        row_info['attributes'][attr_name] = attr_value_text
                
                rows_data.append(row_info)
            
            return JsonResponse({
                'success': True,
                'user_name': user.name,
                'total_rows': rows.count(),
                'rows_data': rows_data,
                'message': '사용자 데이터를 성공적으로 조회했습니다.'
            })
        
    except Exception as e:
        print(f"오류 발생: {str(e)}")
        return JsonResponse({
            'success': False,
            'error': str(e)
        })

def get_company_name_from_row(row):
    """행에서 회사명을 추출하는 헬퍼 함수"""
    try:
        company_attr = row.values.filter(attribute__name='회사명').first()
        if company_attr and company_attr.value:
            return company_attr.value
        return f"회사_{row.id}"
    except:
        return f"회사_{row.id}"

# ==========================
# 파일 텍스트 추출 함수들
# ==========================
def extract_text_from_file(file_path):
    """파일에서 텍스트 추출 - PDF 표 우선 추출, 텍스트 부족시 OCR 자동 fallback"""
    if not file_path or not os.path.exists(file_path):
        return ""
    
    def is_text_extracted_enough(file_path, extracted_text):
        file_size = os.path.getsize(file_path)
        text_length = len(extracted_text)
        if file_size == 0:
            return False
        ratio = text_length / file_size
        return ratio > 0.001  # 0.1% 이상이면 텍스트가 어느 정도 있다고 판단

    try:
        if file_path.endswith(".pdf"):
            import tempfile
            with pdfplumber.open(file_path) as pdf:
                all_tables = []
                for page in pdf.pages:
                    tables = page.extract_tables()
                    for table in tables:
                        all_tables.append(table)
                if all_tables:
                    # 표가 있으면 표만 추출
                    table_texts = []
                    for table in all_tables:
                        table_texts.append('\n'.join(['\t'.join([cell if cell is not None else '' for cell in row]) for row in table if row]))
                    result = '\n\n'.join(table_texts)
                else:
                    # 표가 없으면 기존 방식
                    text = ''
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + '\n'
                    result = text.strip()
            # 텍스트 부족시 OCR 자동 fallback
            if not is_text_extracted_enough(file_path, result):
                ocr_texts = []
                with pdfplumber.open(file_path) as pdf:
                    for i, page in enumerate(pdf.pages):
                        # 페이지를 이미지로 저장
                        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_img:
                            img = page.to_image(resolution=300)
                            img.save(tmp_img.name, format='PNG')
                            # OCR 수행
                            ocr_text = clova_ocr(tmp_img.name, 'jpg')
                            ocr_texts.append(f"[페이지 {i+1} OCR 결과]:\n{ocr_text}")
                            try:
                                os.remove(tmp_img.name)
                            except:
                                pass
                return f"[경고] 이 PDF는 텍스트가 거의 없는 이미지 기반 PDF로 판단되어 OCR로 텍스트를 추출했습니다.\n\n" + '\n\n'.join(ocr_texts)
            return result
        elif file_path.endswith((".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tiff", ".tif", ".webp")):
            return clova_ocr(file_path, "jpg")
        elif file_path.endswith((".hwp", ".hwpx")):
            pdf_path = convert_hwp_to_pdf(file_path)
            if os.path.exists(pdf_path):
                extracted_text = extract_text_from_file(file_path)
                if os.path.exists(pdf_path):
                    os.remove(pdf_path)
                return extracted_text
            else:
                return "HWP 파일 변환 실패"
        elif file_path.endswith((".docx", ".doc")):
            # DOCX 파일 처리
            try:
                from docx import Document
                
                # DOCX 파일 열기
                doc = Document(file_path)
                
                # 모든 단락의 텍스트 추출
                paragraphs = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        paragraphs.append(paragraph.text.strip())
                
                # 모든 테이블의 텍스트 추출
                tables = []
                for table in doc.tables:
                    table_text = []
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            table_text.append(' | '.join(row_text))
                    if table_text:
                        tables.append('\n'.join(table_text))
                
                # 결과 조합
                result = []
                if paragraphs:
                    result.append('\n'.join(paragraphs))
                if tables:
                    result.append('\n\n'.join(tables))
                
                final_result = '\n\n'.join(result) if result else "DOCX 파일에서 텍스트를 추출할 수 없습니다."
                return final_result
                
            except ImportError:
                return "DOCX 파일 처리를 위해 python-docx 라이브러리가 필요합니다."
            except Exception as e:
                return f"DOCX 텍스트 추출 실패: {str(e)}"
        elif file_path.endswith((".txt", ".md", ".markdown", ".rst", ".adoc")):
            # 텍스트 파일과 마크다운 파일 처리
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                    return content.strip()
            except UnicodeDecodeError:
                # UTF-8로 읽기 실패시 다른 인코딩 시도
                try:
                    with open(file_path, 'r', encoding='cp949') as f:
                        content = f.read()
                        return content.strip()
                except UnicodeDecodeError:
                    try:
                        with open(file_path, 'r', encoding='euc-kr') as f:
                            content = f.read()
                            return content.strip()
                    except Exception as e:
                        return f"텍스트 파일 읽기 실패: 인코딩 문제"
            except Exception as e:
                return f"텍스트 파일 읽기 실패: {str(e)}"
        elif file_path.endswith((".csv", ".tsv")):
            # CSV/TSV 파일 처리
            try:
                import csv
                delimiter = ',' if file_path.endswith('.csv') else '\t'
                with open(file_path, 'r', encoding='utf-8') as f:
                    reader = csv.reader(f, delimiter=delimiter)
                    rows = list(reader)
                    if rows:
                        # 헤더와 데이터를 텍스트로 변환
                        result = []
                        for i, row in enumerate(rows):
                            if i == 0:  # 헤더
                                result.append(f"헤더: {' | '.join(row)}")
                            else:  # 데이터
                                result.append(f"행 {i}: {' | '.join(row)}")
                        return '\n'.join(result)
                    else:
                        return "빈 CSV/TSV 파일"
            except Exception as e:
                return f"CSV/TSV 파일 읽기 실패: {str(e)}"
        elif file_path.endswith((".json", ".xml", ".yaml", ".yml")):
            # 구조화된 데이터 파일 처리
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                    return f"구조화된 데이터 파일 내용:\n{content.strip()}"
            except Exception as e:
                return f"구조화된 데이터 파일 읽기 실패: {str(e)}"
        return ""
    except Exception as e:
        return f"텍스트 추출 실패: {str(e)}"

def clova_ocr(file_path, fmt):
    """Clova OCR을 사용한 텍스트 추출"""
    from config import NAVER_CLOVA_OCR_API_KEY, NAVER_CLOUD_CLOVA_OCR_API_URL
    
    request_json = {
        'images': [{'format': fmt, 'name': 'demo'}],
        'requestId': str(uuid.uuid4()),
        'version': 'V1',
        'timestamp': int(time.time() * 1000)
    }
    payload = {'message': json.dumps(request_json).encode('UTF-8')}
    
    try:
        with open(file_path, 'rb') as f:
            files = [('file', f)]
            headers = {'X-OCR-SECRET': NAVER_CLOVA_OCR_API_KEY}
            
            response = requests.post(NAVER_CLOUD_CLOVA_OCR_API_URL, headers=headers, data=payload, files=files)
            
            if response.status_code != 200:
                return ""
            
            response_data = response.json()
            if 'images' not in response_data or not response_data['images']:
                return ""
            
            full_text = ""
            for field in response_data['images'][0].get('fields', []):
                full_text += field['inferText'] + " "
            
            extracted_text = full_text.strip()
            return extracted_text
            
    except Exception as e:
        return ""

def check_libreoffice_status():
    """LibreOffice 설치 상태와 버전을 확인합니다."""
    try:
        result = subprocess.run([
            "libreoffice", "--version"
        ], stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=30)
        
        if result.returncode == 0:
            return True
        else:
            return False
    except Exception as e:
        return False

def convert_hwp_to_pdf(hwp_path):
    """HWP를 PDF로 변환"""
    output_dir = os.path.dirname(hwp_path)
    try:
        # 파일 크기 확인
        file_size = os.path.getsize(hwp_path)
        
        # 파일 크기에 따른 timeout 조정
        if file_size > 50 * 1024 * 1024:  # 50MB 이상
            timeout = 1800  # 30분
        elif file_size > 10 * 1024 * 1024:  # 10MB 이상
            timeout = 900   # 15분
        else:
            timeout = 600   # 10분 (기본값)
        
        result = subprocess.run([
            "libreoffice",
            "--headless",
            "--convert-to", "pdf:writer_pdf_Export",
            hwp_path,
            "--outdir", output_dir
        ], stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=timeout)

        basename = os.path.splitext(os.path.basename(hwp_path))[0] + ".pdf"
        converted_pdf = os.path.join(output_dir, basename)

        if os.path.exists(converted_pdf):
            return converted_pdf
        else:
            return ""
            
    except subprocess.TimeoutExpired:
        # LibreOffice 프로세스 강제 종료
        try:
            subprocess.run(["pkill", "-f", "libreoffice"], timeout=10)
        except:
            pass
        return ""
    except Exception as e:
        return ""

def download_file_from_s3_key(s3_key):
    """S3 키를 사용하여 파일 다운로드"""
    try:
        from config import AWS_S3_ACCESS_KEY, AWS_S3_SECRET_KEY, AWS_S3_BUCKET_NAME, AWS_S3_REGION
        
        # S3 클라이언트 생성
        s3_client = boto3.client(
            's3',
            aws_access_key_id=AWS_S3_ACCESS_KEY,
            aws_secret_access_key=AWS_S3_SECRET_KEY,
            region_name=AWS_S3_REGION
        )
        
        # 임시 파일 생성
        temp_dir = tempfile.gettempdir()
        file_name = s3_key.split('/')[-1]
        temp_path = os.path.join(temp_dir, f"s3_{file_name}")
        
        # S3에서 파일 다운로드
        s3_client.download_file(AWS_S3_BUCKET_NAME, s3_key, temp_path)
        
        return temp_path
        
    except Exception as e:
        return None

def download_file_from_url(url):
    """URL에서 파일 다운로드"""
    try:
        # S3 URL인 경우 boto3 사용
        if 's3.ap-northeast-2.amazonaws.com' in url:
            return download_file_from_s3(url)
        
        # 일반 URL인 경우 requests 사용
        temp_dir = tempfile.gettempdir()
        file_name = url.split('/')[-1].split('?')[0]  # 쿼리 파라미터 제거
        temp_path = os.path.join(temp_dir, file_name)
        
        # 파일 다운로드
        response = requests.get(url, stream=True, timeout=30)
        response.raise_for_status()
        
        with open(temp_path, 'wb') as f:
            for chunk in response.iter_content(chunk_size=8192):
                f.write(chunk)
        
        return temp_path
    except Exception as e:
        return None

def download_file_from_s3(url):
    """S3에서 직접 파일 다운로드"""
    try:
        from config import AWS_S3_ACCESS_KEY, AWS_S3_SECRET_KEY, AWS_S3_BUCKET_NAME, AWS_S3_REGION
        
        # S3 클라이언트 생성
        s3_client = boto3.client(
            's3',
            aws_access_key_id=AWS_S3_ACCESS_KEY,
            aws_secret_access_key=AWS_S3_SECRET_KEY,
            region_name=AWS_S3_REGION
        )
        
        # URL에서 S3 키 추출
        bucket_name = AWS_S3_BUCKET_NAME
        
        # URL에서 파일 경로 추출
        if '/media/' in url:
            s3_key = 'media/' + url.split('/media/')[-1].split('?')[0]
        elif '/note_files/' in url:
            s3_key = 'note_files/' + url.split('/note_files/')[-1].split('?')[0]
        else:
            # 다른 경로인 경우 전체 경로에서 추출
            s3_key = url.split(f'{bucket_name}/')[-1].split('?')[0]
        
        # 임시 파일 생성
        temp_dir = tempfile.gettempdir()
        file_name = s3_key.split('/')[-1]
        temp_path = os.path.join(temp_dir, file_name)
        
        # S3에서 파일 다운로드
        s3_client.download_file(bucket_name, s3_key, temp_path)
        
        return temp_path
    except Exception as e:
        return None

def extract_file_text(field_name, file_info):
    """파일 정보에서 텍스트 추출하는 헬퍼 함수"""
    try:
        # 파일 크기 체크 (50MB 제한)
        file_size = file_info.get('file_size', 0)
        if file_size > 50 * 1024 * 1024:  # 50MB
            return f"[{field_name} - {file_info.get('original_filename', '파일')}]: 파일이 너무 커서 텍스트 추출을 건너뜁니다.", None
        
        file_path = None
        # S3 키가 있으면 직접 사용
        s3_key = file_info.get('s3_key')
        if s3_key:
            file_path = download_file_from_s3_key(s3_key)
        # S3 키가 없고 download_url이 있으면 사용
        elif file_info.get('download_url'):
            file_path = download_file_from_url(file_info['download_url'])
        
        if file_path and os.path.exists(file_path):
            file_text = extract_text_from_file(file_path)
            file_hash = hashlib.md5(open(file_path, 'rb').read()).hexdigest()
            
            if file_text:
                result = f"[{field_name} - {file_info.get('original_filename', '파일')}]:\n{file_text}"
                return result, file_hash
            else:
                result = f"[{field_name} - {file_info.get('original_filename', '파일')}]: 텍스트 추출 실패"
                return result, file_hash
            
            # 임시 파일 정리
            try:
                os.remove(file_path)
            except:
                pass
            
            return result, file_hash
        else:
            return f"[{field_name} - {file_info.get('original_filename', '파일')}]: 파일을 다운로드할 수 없습니다.", None
    except Exception as e:
        return f"[{field_name} - {file_info.get('original_filename', '파일')}]: 텍스트 추출 중 오류 발생", None

def get_row_files_text(row):
    """행의 모든 파일에서 텍스트 추출"""
    file_texts = []
    
    try:
        # 파일 속성값들 처리
        file_attribute_values = AttributeValue.objects.filter(row=row, attribute__attributeType__name='file')
        
        for attr_value in file_attribute_values:
            attr_name = attr_value.attribute.name
            
            if attr_value.value:
                try:
                    # 파일 데이터 파싱
                    file_data = json.loads(attr_value.value) if isinstance(attr_value.value, str) else attr_value.value
                    
                    # 음성파일 속성인 경우 (data 구조)
                    if isinstance(file_data, dict) and 'data' in file_data:
                        for file_id, file_info in file_data['data'].items():
                            if file_info.get('type') in ['file', 'image']:
                                # text 타입은 직접 처리
                                if file_info.get('type') == 'text':
                                    text_content = file_info.get('text', '')
                                    if text_content:
                                        file_texts.append(f"[{attr_name} - 텍스트]:\n{text_content}")
                                else:
                                    # 파일에서 텍스트 추출
                                    file_text, file_hash = extract_file_text(attr_name, file_info)
                                    if file_text:
                                        file_texts.append(file_text)
                    
                    # 일반 파일 속성인 경우 (배열 구조)
                    elif isinstance(file_data, list):
                        for file_info in file_data:
                            if file_info.get('type') == 'text':
                                text_content = file_info.get('text', '')
                                if text_content:
                                    file_texts.append(f"[{attr_name} - 텍스트]:\n{text_content}")
                            else:
                                # 파일에서 텍스트 추출
                                file_text, file_hash = extract_file_text(attr_name, file_info)
                                if file_text:
                                    file_texts.append(file_text)
                    
                    # 단일 파일 경로인 경우
                    elif isinstance(file_data, str):
                        file_path = file_data
                        if file_path.startswith('http'):
                            file_path = download_file_from_url(file_path)
                        
                        if file_path and os.path.exists(file_path):
                            file_text = extract_text_from_file(file_path)
                            if file_text:
                                file_texts.append(f"[{attr_name} 파일 내용]:\n{file_text}")
                            
                            # 임시 파일 정리
                            if file_data.startswith('http'):
                                try:
                                    os.remove(file_path)
                                except:
                                    pass
                                    
                except Exception as e:
                    print(f"파일 처리 중 오류: {e}")
                    
    except Exception as e:
        print(f"행 파일 텍스트 추출 실패: {e}")
    
    return file_texts

# ==========================
# OpenAI API 연동 함수들
# ==========================
def generate_business_overview_with_openai(row_data, service_product, row=None, file_texts=None):
    """OpenAI API를 통해 사업 개요 추천 생성"""
    try:
        if not OPEN_AI_API_KEY:
            print("OpenAI API 키가 설정되지 않았습니다.")
            return None
        
        # 파일 텍스트가 없으면 행에서 추출
        if file_texts is None and row:
            file_texts = get_row_files_text(row)
            print(f"행에서 파일 텍스트 추출: {len(file_texts)}개")
        
        # 기업 정보를 OpenAI에 전달할 프롬프트 구성
        company_info = format_company_info_for_openai(row_data, file_texts)
        print(f"기업 정보: {company_info}")
        
        # 기업 정보에서 매출액 추출
        revenue = row_data.get("매출액", row_data.get("매출액(백만원)", 0))
        
        system_prompt = """당신은 사업계획서 작성 전문가입니다. 
        제공된 기업 정보와 사용자가 입력한 주 서비스·생산품목을 바탕으로 사업 개요를 작성해주세요(서비스 및 상품의 주요 내용, 제품의 다양성, 인지도 등).
        
        **중요한 지침:**
        1. 제공된 기업 정보와 사용자 입력 정보를 기반으로 답변하세요.
        2. 첨부된 파일 내용이 있으면 이를 참고하여 더 구체적이고 정확한 사업 개요를 작성해주세요.
        3. 사업 개요는 구체적이고 현실적으로 작성해주세요.
        4. 해당 서비스/제품의 용도와 특성을 명확히 설명해주세요.
        5. 200-300자 내외로 작성해주세요.
        
        다음 JSON 형식으로 응답해주세요:
        {
            "business_overview": "사업의 주요 내용과 특징, 서비스/제품의 용도 및 특성을 포함한 구체적인 설명"
        }"""
        
        # 파일 텍스트가 있으면 포함
        file_context = ""
        if file_texts:
            file_context = "\n\n**첨부 파일 내용:**\n" + "\n\n".join(file_texts)
        
        user_message = f"""다음 기업 정보와 사용자 입력 정보를 바탕으로 사업 개요를 작성해주세요:

**사용자 입력 정보:**
- 주 서비스·생산품목: {service_product}

**기업 정보:**
{company_info}{file_context}

위 정보를 바탕으로 사업 개요를 작성해주세요."""
        
        # OpenAI API 호출
        headers = {
            'Authorization': f'Bearer {OPEN_AI_API_KEY}',
            'Content-Type': 'application/json'
        }
        
        payload = {
            'model': 'gpt-4',
            'messages': [
                {'role': 'system', 'content': system_prompt},
                {'role': 'user', 'content': user_message}
            ],
            'max_tokens': 1000,
            'temperature': 0.3
        }
        
        response = requests.post(
            'https://api.openai.com/v1/chat/completions',
            headers=headers,
            json=payload,
            timeout=60
        )
        
        if response.status_code == 200:
            result = response.json()
            ai_response = result['choices'][0]['message']['content'].strip()
            
            # JSON 응답 파싱
            try:
                # JSON 블록 추출 (```json ... ``` 형태일 수 있음)
                if '```json' in ai_response:
                    json_start = ai_response.find('```json') + 7
                    json_end = ai_response.find('```', json_start)
                    if json_end != -1:
                        json_str = ai_response[json_start:json_end].strip()
                    else:
                        json_str = ai_response[json_start:].strip()
                else:
                    json_str = ai_response
                
                recommendation_data = json.loads(json_str)
                print(f"OpenAI 추천 응답 파싱 성공: {recommendation_data}")
                return recommendation_data
                
            except json.JSONDecodeError as e:
                print(f"JSON 파싱 실패: {e}")
                print(f"AI 응답: {ai_response}")
                return None
        else:
            print(f"OpenAI API 호출 실패: {response.status_code}")
            return None
            
    except Exception as e:
        print(f"OpenAI API 연동 중 오류: {e}")
        return None

def generate_business_plan_with_openai(row_data):
    """OpenAI API를 통해 사업계획서 데이터 생성"""
    try:
        if not OPEN_AI_API_KEY:
            print("OpenAI API 키가 설정되지 않았습니다.")
            return None
        
        # 기업 정보를 OpenAI에 전달할 프롬프트 구성
        company_info = format_company_info_for_openai(row_data)
        
        system_prompt = """당신은 사업계획서 작성 전문가입니다. 
        제공된 기업 정보를 바탕으로 신용취약 소상공인 자금 사업계획서에 필요한 모든 정보를 생성해주세요.
        
        **중요한 지침:**
        1. 제공된 기업 정보만을 기반으로 답변하세요.
        2. 정보가 부족한 경우 "정보 부족"이라고 명시하세요.
        3. 추측하지 말고 확실한 정보만 사용하세요.
        4. 모든 값은 현실적이고 구체적으로 작성해주세요.
        
        다음 JSON 형식으로 응답해주세요:
        {
            "주 서비스·생산품목": "기업의 주요 서비스나 생산품목 (정보 부족시 '정보 부족')",
            "매출액(백만원)": 숫자값 (정보 부족시 0),
            "주 사용 플랫폼": "주로 사용하는 플랫폼이나 채널 (정보 부족시 '정보 부족')",
            "점포 보유현황": "유" 또는 "무" (정보 부족시 '정보 부족'),
            "대표자_경력": [
                ["기간", "근무처", "담당업무", "최종직위"]
            ],
            "실제경영자_관계": "대표자와의 관계 (정보 부족시 '본인')",
            "실제경영자_경력": [
                ["기간", "근무처", "담당업무", "최종직위"]
            ],
            "사업개요": "사업의 주요 내용과 특징 (정보 부족시 '정보 부족')",
            "향후 사업계획(자금용도 포함)": "자금을 어떻게 사용할지 포함한 향후 계획 (정보 부족시 '정보 부족')",
            "자금소요_운전자금": 숫자값 (정보 부족시 30),
            "자금조달_본건차입금": 숫자값 (정보 부족시 20),
            "자금조달_자체자금": 숫자값 (정보 부족시 10),
            "자금조달_기타": 숫자값 (정보 부족시 0)
        }
        
        **경력 데이터 처리:**
        - 경력 정보가 없으면 빈 배열 [] 반환
        - 경력 정보가 있으면 ["기간", "근무처", "담당업무", "최종직위"] 형식으로 반환
        - 각 항목이 부족하면 빈 문자열 "" 사용"""
        
        user_message = f"다음 기업 정보를 바탕으로 사업계획서 데이터를 생성해주세요:\n\n{company_info}"
        
        # OpenAI API 호출
        headers = {
            'Authorization': f'Bearer {OPEN_AI_API_KEY}',
            'Content-Type': 'application/json'
        }
        
        payload = {
            'model': 'gpt-4',
            'messages': [
                {'role': 'system', 'content': system_prompt},
                {'role': 'user', 'content': user_message}
            ],
            'max_tokens': 2000,
            'temperature': 0.3
        }
        
        response = requests.post(
            'https://api.openai.com/v1/chat/completions',
            headers=headers,
            json=payload,
            timeout=60
        )
        
        if response.status_code == 200:
            result = response.json()
            ai_response = result['choices'][0]['message']['content'].strip()
            
            # JSON 응답 파싱
            try:
                # JSON 블록 추출 (```json ... ``` 형태일 수 있음)
                if '```json' in ai_response:
                    json_start = ai_response.find('```json') + 7
                    json_end = ai_response.find('```', json_start)
                    if json_end != -1:
                        json_str = ai_response[json_start:json_end].strip()
                    else:
                        json_str = ai_response[json_start:].strip()
                else:
                    json_str = ai_response
                
                business_plan_data = json.loads(json_str)
                print(f"OpenAI 응답 파싱 성공: {business_plan_data}")
                return business_plan_data
                
            except json.JSONDecodeError as e:
                print(f"JSON 파싱 실패: {e}")
                print(f"AI 응답: {ai_response}")
                return None
        else:
            print(f"OpenAI API 호출 실패: {response.status_code}")
            return None
            
    except Exception as e:
        print(f"OpenAI API 연동 중 오류: {e}")
        return None

def format_company_info_for_openai(row_data, file_texts=None):
    """기업 정보를 OpenAI에 전달하기 위한 형식으로 변환"""
    formatted_info = []
    
    for key, value in row_data.items():
        if value and str(value).strip() and key != "파일_텍스트":  # 파일_텍스트는 별도 처리
            formatted_info.append(f"{key}: {value}")
    
    # 파일 텍스트가 있으면 추가
    if file_texts and isinstance(file_texts, list):
        formatted_info.append("\n첨부 파일 내용:")
        for file_text in file_texts:
            formatted_info.append(file_text)
    
    return "\n".join(formatted_info)

# ==========================
# DOCX 조작 유틸리티 함수들
# ==========================
def cell_right(table, r, c):
    """테이블에서 오른쪽 셀 반환"""
    try:
        if table is None or r < 0 or c < 0:
            return None
        
        # 행 범위 체크
        if r >= len(table.rows):
            return None
        
        row = table.rows[r]
        # 열 범위 체크
        if c + 1 >= len(row.cells):
            return None
        
        return row.cells[c + 1]
    except Exception as e:
        print(f"cell_right 오류 (r={r}, c={c}): {e}")
        return None

def cell_below(table, r, c):
    """테이블에서 아래쪽 셀 반환"""
    try:
        if table is None or r < 0 or c < 0:
            return None
        
        # 행 범위 체크
        if r + 1 >= len(table.rows):
            return None
        
        row = table.rows[r + 1]
        # 열 범위 체크
        if c >= len(row.cells):
            return None
        
        return row.cells[c]
    except Exception as e:
        print(f"cell_below 오류 (r={r}, c={c}): {e}")
        return None

def set_cell_text(cell, value):
    """셀에 텍스트 설정"""
    try:
        if cell is None:
            print("셀이 None입니다.")
            return
        
        # 기존 문단 제거하고 한 문단으로 교체
        try:
            for p in cell.paragraphs:
                if p:
                    p.clear()
        except Exception as clear_error:
            print(f"문단 정리 중 오류: {clear_error}")
        
        # 새 텍스트 설정
        try:
            cell.text = str(value)
        except Exception as text_error:
            print(f"텍스트 설정 중 오류: {text_error}")
            # 대안: 첫 번째 문단에 텍스트 추가
            try:
                if cell.paragraphs and len(cell.paragraphs) > 0:
                    cell.paragraphs[0].text = str(value)
                else:
                    # 문단이 없으면 새로 생성
                    cell.add_paragraph(str(value))
            except Exception as alt_error:
                print(f"대안 텍스트 설정도 실패: {alt_error}")
                
    except Exception as e:
        print(f"set_cell_text 전체 오류: {e}")
        # 최후의 수단: 아무것도 하지 않음
        pass

def norm(s):
    """텍스트 정규화"""
    if not s:
        return ""
    return re.sub(r"\s+", "", str(s)).strip()

def is_match(label, cell_text, aliases=None):
    """라벨 매칭 확인"""
    if aliases is None:
        aliases = []
    t = norm(cell_text)
    if norm(label) in t:
        return True
    return any(norm(a) in t for a in aliases)

# 라벨 별 동의어
ALIASES = {
    "주 서비스·생산품목": ["주서비스·생산품목", "주서비스/생산품목", "주서비스", "생산품목"],
    "매출액": ["매출액(백만원)", "매출액백만원"],
    "주 사용 플랫폼": ["주사용플랫폼", "사용플랫폼"],
    "점포 보유현황": ["점포보유현황", "점포보유"],
    "사업개요": ["주서비스·생산품목의용도및특성", "서비스및상품의주요내용"],
    "향후 사업계획": ["자금용도및사업계획", "향후사업계획", "자금용도"],
    "자금소요": ["자금소요내역", "자금소요(백만원)"],
    "자금조달": ["자금조달계획", "자금조달(백만원)"],
    "경력": ["경력", "대표자경력", "실제경영자경력"],
}

def find_and_fill_simple_labels(doc, fill_map):
    """표/문단에서 라벨을 찾아 값을 채우는 함수"""
    try:
        # 1) 표 내부
        for table in doc.tables:
            try:
                rows = len(table.rows)
                for r in range(rows):
                    try:
                        row = table.rows[r]
                        for c in range(len(row.cells)):
                            try:
                                cell = row.cells[c]
                                txt = cell.text if cell.text else ""
                                if not txt.strip():
                                    continue
                                
                                for label, value in fill_map.items():
                                    if is_match(label, txt, ALIASES.get(label, [])):
                                        try:
                                            # 우선 오른쪽 셀에 써보고, 없으면 아래 셀 시도
                                            target = cell_right(table, r, c) or cell_below(table, r, c)
                                            if target:
                                                set_cell_text(target, value)
                                            else:
                                                # 같은 셀 자체를 치환해야 하는 경우
                                                set_cell_text(cell, value)
                                        except Exception as cell_error:
                                            print(f"셀 채우기 중 오류: {cell_error}")
                                            continue
                            except Exception as col_error:
                                print(f"열 {c} 처리 중 오류: {col_error}")
                                continue
                    except Exception as row_error:
                        print(f"행 {r} 처리 중 오류: {row_error}")
                        continue
            except Exception as table_error:
                print(f"테이블 처리 중 오류: {table_error}")
                continue

        # 2) 문단
        try:
            paragraphs = doc.paragraphs
            for i, p in enumerate(paragraphs):
                try:
                    for label, value in fill_map.items():
                        if is_match(label, p.text, ALIASES.get(label, [])):
                            try:
                                # 다음 문단이 비어있으면 거기에, 아니면 라벨 문단 바로 뒤에 새 문단 삽입
                                target_index = i+1 if i+1 < len(paragraphs) else i
                                if target_index < len(paragraphs) and not paragraphs[target_index].text.strip():
                                    paragraphs[target_index].text = str(value)
                                else:
                                    p.insert_paragraph_before(str(value))
                            except Exception as para_error:
                                print(f"문단 채우기 중 오류: {para_error}")
                                continue
                except Exception as para_iter_error:
                    print(f"문단 {i} 처리 중 오류: {para_iter_error}")
                    continue
        except Exception as para_section_error:
            print(f"문단 섹션 처리 중 오류: {para_section_error}")
            
    except Exception as e:
        print(f"find_and_fill_simple_labels 전체 오류: {e}")
        raise

def fill_shop_checkbox_format(text, has_shop):
    """점포 보유현황 체크박스 형식 처리"""
    has = norm(has_shop) in ["유", "있음", "예", "true"]
    # 원문이 '☑ 유 □ 무'나 '□ 유 □ 무' 등일 수 있어 모두 정규화 처리
    base = re.sub(r"[☑■□]\s*유", "□ 유", text)
    base = re.sub(r"[☑■□]\s*무", "□ 무", base)
    if has:
        base = base.replace("□ 유", "☑ 유")
    else:
        base = base.replace("□ 무", "☑ 무")
    return base

def apply_shop_checkbox(doc, has_shop):
    """점포 보유현황 체크박스 적용"""
    try:
        for table in doc.tables:
            try:
                for row in table.rows:
                    try:
                        for cell in row.cells:
                            try:
                                if cell and cell.text:
                                    t = cell.text
                                    if is_match("점포 보유현황", t, ALIASES["점포 보유현황"]) and ("유" in t and "무" in t):
                                        formatted_text = fill_shop_checkbox_format(t, has_shop)
                                        set_cell_text(cell, formatted_text)
                                        print(f"점포 보유현황 체크박스 적용: {has_shop}")
                                        return  # 첫 번째 매칭되는 셀만 처리
                            except Exception as cell_error:
                                print(f"점포 보유현황 셀 처리 중 오류: {cell_error}")
                                continue
                    except Exception as row_error:
                        print(f"점포 보유현황 행 처리 중 오류: {row_error}")
                        continue
            except Exception as table_error:
                print(f"점포 보유현황 테이블 처리 중 오류: {table_error}")
                continue
    except Exception as e:
        print(f"apply_shop_checkbox 전체 오류: {e}")
        # 오류가 발생해도 계속 진행
        pass

def fill_career_table(doc, label_text, careers):
    """경력 표 채우기"""
    try:
        for table in doc.tables:
            try:
                # 테이블 텍스트 수집 (안전하게)
                table_text_parts = []
                for row in table.rows:
                    for cell in row.cells:
                        if cell and cell.text:
                            table_text_parts.append(cell.text)
                
                table_text = "\n".join(table_text_parts)
                
                if is_match(label_text, table_text, ALIASES["경력"]) and \
                   ("구 분" in table_text or "구분" in table_text) and "근무처" in table_text:
                    
                    print(f"경력 표 발견: {label_text}")
                    
                    # 헤더 다음 행부터 입력
                    needed_rows = len(careers) if careers else 0
                    header_row_idx = 0
                    start_row_idx = header_row_idx + 1
                    
                    # 필요한 행 수만큼 추가
                    while len(table.rows) < start_row_idx + needed_rows:
                        try:
                            table.add_row()
                        except Exception as add_row_error:
                            print(f"행 추가 중 오류: {add_row_error}")
                            break

                    # 각 행 채우기 (안전한 인덱스 접근)
                    for i, career in enumerate(careers):
                        try:
                            r = start_row_idx + i
                            
                            # 행 범위 체크
                            if r >= len(table.rows):
                                print(f"행 {r}가 테이블 범위를 벗어남")
                                break
                            
                            # career가 리스트인지 확인하고 안전하게 접근
                            if isinstance(career, list) and len(career) >= 4:
                                period = career[0] if career[0] else ""
                                company = career[1] if career[1] else ""
                                duty = career[2] if career[2] else ""
                                title = career[3] if career[3] else ""
                            else:
                                # career가 리스트가 아니거나 길이가 부족한 경우
                                period = str(career) if career else ""
                                company = ""
                                duty = ""
                                title = ""
                            
                            # 안전한 셀 접근
                            row = table.rows[r]
                            if len(row.cells) > 1:
                                set_cell_text(row.cells[1], period)
                            if len(row.cells) > 2:
                                set_cell_text(row.cells[2], company)
                            if len(row.cells) > 3:
                                set_cell_text(row.cells[3], duty)
                            if len(row.cells) > 4:
                                set_cell_text(row.cells[4], title)
                                
                        except Exception as e:
                            print(f"경력 표 채우기 중 오류 (행 {i}): {e}")
                            continue
                    
                    print(f"경력 표 채우기 완료: {label_text}")
                    return  # 첫 번째 매칭되는 테이블만 처리
                    
            except Exception as table_error:
                print(f"테이블 처리 중 오류: {table_error}")
                continue
                
    except Exception as e:
        print(f"fill_career_table 전체 오류: {e}")
        raise

def fill_money_tables(doc, data):
    """자금소요/자금조달 표 채우기"""
    try:
        need = data.get("자금소요_운전자금", 0)
        need_total = data.get("자금소요_합계", need)

        loan = data.get("자금조달_본건차입금", 0)
        self_ = data.get("자금조달_자체자금", 0)
        other = data.get("자금조달_기타", 0)
        fund_total = data.get("자금조달_합계", loan + self_ + other)

        def fmt(v):
            try:
                return str(int(v))
            except:
                return str(v)

        for table in doc.tables:
            try:
                # 테이블 텍스트 수집 (안전하게)
                table_text_parts = []
                for row in table.rows:
                    for cell in row.cells:
                        if cell and cell.text:
                            table_text_parts.append(cell.text)
                
                table_text = "\n".join(table_text_parts)

                # 자금소요 표
                if any(k in table_text for k in ["자금소요", "자금소요내역"]):
                    print("자금소요 표 발견")
                    for r, row in enumerate(table.rows):
                        try:
                            row_text_parts = []
                            for c in row.cells:
                                if c and c.text:
                                    row_text_parts.append(c.text)
                            
                            row_text = " ".join(row_text_parts)
                            
                            if "운전" in row_text and len(row.cells) > 0:
                                set_cell_text(row.cells[-1], fmt(need))
                            if "합계" in row_text and len(row.cells) > 0:
                                set_cell_text(row.cells[-1], fmt(need_total))
                        except Exception as row_error:
                            print(f"자금소요 행 {r} 처리 중 오류: {row_error}")
                            continue

                # 자금조달 표
                if any(k in table_text for k in ["자금조달", "자금조달계획"]):
                    print("자금조달 표 발견")
                    for r, row in enumerate(table.rows):
                        try:
                            row_text_parts = []
                            for c in row.cells:
                                if c and c.text:
                                    row_text_parts.append(c.text)
                            
                            row_text = " ".join(row_text_parts)
                            
                            if "본건" in row_text and len(row.cells) > 0:
                                set_cell_text(row.cells[-1], fmt(loan))
                            if "자체" in row_text and len(row.cells) > 0:
                                set_cell_text(row.cells[-1], fmt(self_))
                            if "기타" in row_text or "은행차입금등 기타" in row_text:
                                if len(row.cells) > 0:
                                    set_cell_text(row.cells[-1], fmt(other))
                            if "합 계" in row_text or "합계" in row_text:
                                if len(row.cells) > 0:
                                    set_cell_text(row.cells[-1], fmt(fund_total))
                        except Exception as row_error:
                            print(f"자금조달 행 {r} 처리 중 오류: {row_error}")
                            continue
                            
            except Exception as table_error:
                print(f"자금 표 처리 중 오류: {table_error}")
                continue
                
    except Exception as e:
        print(f"fill_money_tables 전체 오류: {e}")
        raise

def fill_business_plan_docx(doc, business_plan_data):
    """사업계획서 DOCX 파일 채우기 메인 함수"""
    try:
        # 1) 간단 라벨-값 매핑
        simple_map = {
            "주 서비스·생산품목": business_plan_data.get("주 서비스·생산품목", "정보 부족"),
            "매출액": f'{business_plan_data.get("매출액(백만원)", 0)} 백만원',
            "주 사용 플랫폼": business_plan_data.get("주 사용 플랫폼", "정보 부족"),
            "사업개요": business_plan_data.get("사업개요", "정보 부족"),
            "향후 사업계획": business_plan_data.get("향후 사업계획(자금용도 포함)", "정보 부족"),
        }
        find_and_fill_simple_labels(doc, simple_map)
        print("1) 라벨-값 매핑 완료")

        # 2) 점포 보유현황 체크박스 처리
        try:
            apply_shop_checkbox(doc, business_plan_data.get("점포 보유현황", "정보 부족"))
            print("2) 점포 보유현황 체크박스 처리 완료")
        except Exception as e:
            print(f"점포 보유현황 처리 중 오류: {e}")

        # 3) 경력 표 채우기
        try:
            fill_career_table(doc, "대표자", business_plan_data.get("대표자_경력", []))
            print("3-1) 대표자 경력 표 채우기 완료")
        except Exception as e:
            print(f"대표자 경력 표 채우기 중 오류: {e}")
        
        try:
            find_and_fill_simple_labels(doc, {"실제경영자": f"(대표자와의 관계 : {business_plan_data.get('실제경영자_관계','본인')})"})
            print("3-2) 실제경영자 관계 처리 완료")
        except Exception as e:
            print(f"실제경영자 관계 처리 중 오류: {e}")
        
        try:
            fill_career_table(doc, "실제경영자", business_plan_data.get("실제경영자_경력", []))
            print("3-3) 실제경영자 경력 표 채우기 완료")
        except Exception as e:
            print(f"실제경영자 경력 표 채우기 중 오류: {e}")

        # 4) 자금 소요/조달 표 채우기
        try:
            fill_money_tables(doc, business_plan_data)
            print("4) 자금 소요/조달 표 채우기 완료")
        except Exception as e:
            print(f"자금 소요/조달 표 채우기 중 오류: {e}")
        
        print("사업계획서 DOCX 채우기 완료")
        
    except Exception as e:
        print(f"사업계획서 DOCX 채우기 중 전체 오류: {e}")
        raise